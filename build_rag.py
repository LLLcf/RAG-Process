import os
import re
import json
import uuid
import pickle
import warnings
import requests # 新增：用于 Reranker API 调用
from typing import List, Dict, Any, Union
from dataclasses import dataclass

import numpy as np
import torch
import faiss
import jieba
import jieba.analyse
import docx
from rank_bm25 import BM25Okapi

# 本地模型相关库 (仅在 local 模式下需要，但也保留引用防止报错)
from transformers import AutoModel, AutoTokenizer, AutoModelForCausalLM
from vllm import LLM, SamplingParams

# LlamaIndex 核心组件
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import NodeWithScore
from llama_index.core.llms import ChatMessage

# OpenAI SDK (用于 API 模式)
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None # 容错处理

# 忽略警告
warnings.filterwarnings('ignore')

# ================= 配置类 =================

@dataclass
class EnhancedConfig:
    """增强配置类"""
    
    # --- 运行模式选择 ---
    # 可选: "local" (使用本地显卡加载模型) 或 "api" (调用远程接口)
    MODE = "local" 
    
    # === 本地模型路径 (MODE="local" 时生效) ===
    EMBEDDING_MODEL = "/root/lanyun-fs/models/Qwen3-Embedding-0.6B"
    RERANKER_MODEL = "/root/lanyun-fs/models/Qwen3-Reranker-0.6B"
    GENERATION_MODEL = "/root/lanyun-tmp/models/Qwen3-4B"
    
    # === API 配置 (MODE="api" 时生效) ===
    # 1. LLM API (兼容 OpenAI 格式, 如 DeepSeek, Moonshot, ChatGPT)
    API_BASE_URL = "https://api.deepseek.com/v1" 
    API_KEY = "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"
    API_LLM_MODEL_NAME = "deepseek-chat"
    
    # 2. Embedding API
    API_EMBED_BASE_URL = "https://api.openai.com/v1" # 或其他兼容地址
    API_EMBED_KEY = "sk-xxxxxxxx"
    API_EMBED_MODEL_NAME = "text-embedding-3-small"
    
    # 3. Reranker API (通用 HTTP 接口，如 SiliconFlow, Jina)
    API_RERANK_URL = "https://api.siliconflow.cn/v1/rerank"
    API_RERANK_KEY = "sk-xxxxxxxx"
    API_RERANK_MODEL_NAME = "BAAI/bge-reranker-v2-m3"

    # --- 基础参数 ---
    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
    CHUNK_SIZE = 512  
    CHUNK_OVERLAP = 128
    
    # --- 缓存路径 ---
    CACHE_DIR = "../data/FAISS_Vector_DB"
    NODES_CACHE_FILE = "nodes_cache.pkl"
    DOC_STORE_CACHE_FILE = "doc_store_cache.pkl"
    
    # --- 检索参数 ---
    SIMILARITY_TOP_K = 30  
    BM25_TOP_K = 30
    RERANK_TOP_N = 5       
    
    # --- 生成参数 ---
    MAX_NEW_TOKENS = 4096
    TEMPERATURE = 0.1 
    TOP_P = 0.95
    # 注意：本地 vllm 参数
    gpu_memory_utilization = 0.8
    max_model_len = 32000
    
    # 控制上下文最大长度
    MAX_CONTEXT_LENGTH = 25000 
    
    # --- 功能开关 ---
    QUERY_REWRITE_ENABLED = False
    BM25_ENABLED = True
    RERANKER_ENABLED = True
    RRF_ENABLED = True
    RRF_K = 60

    # --- 输出字段定义 ---
    OUTPUT_COLUMNS = [
        '姓名（中文）', '姓名（原文）', '国籍', '民族/种族', '语种', 
        '出生年份(如1950年)', '出生地', '身体状况（近年）', 
        '任职信息', '教育信息', '家庭关系（如“关系：人名”）', '社会关系（如“关系：人名”）', 
        '社交网络', '社会影响', '个人重要成就/著作', '关键事件', '重要活动', 
        '兴趣偏好', '性格类型倾向', '性格弱点', '价值取向', 
        '职业路径分析', '未来发展预期'
    ]
    
    # --- 提取字段配置 ---
    EXTRACT_GROUPS = {
        "基础身份": [
            '姓名（原文）', '国籍', '民族/种族', '语种', 
            '出生年份(如1950年)', '出生地', '身体状况（近年）', 
            '家庭关系（如“关系：人名”）', '社会关系（如“关系：人名”）'
        ],
        "生涯成就": [
            '任职信息', '教育信息', '社交网络', '社会影响', 
            '个人重要成就/著作', '关键事件', '重要活动'
        ],
        "深度画像": [
            '兴趣偏好', '性格类型倾向', '性格弱点', '价值取向', 
            '职业路径分析', '未来发展预期'
        ]
    }

    # --- Few-Shot 范例数据 ---
    FEW_SHOT_EXAMPLES = {
        "基础身份": """
{
  "姓名（原文）": "Daniel Newham",
  "国籍": "英国",
  "民族/种族": "英德两国血统",
  "语种": "英语、汉语、德语、法语",
  "出生年份(如1950年)": "1980年",
  "出生地": "英国切尔滕纳姆镇",
  "身体状况（近年）": "近年公开信息显示体态稳健，无显著健康问题",
  "家庭关系（如“关系：人名”）": "未知",
  "社会关系（如“关系：人名”）": "同事：克莱奥·吕登"
}
""",
        "生涯成就": """
{
  "任职信息": "现任职务：大牛（Daniel Newham），法学硕士，职业主持人，清华大学艺术学院博士\\n曾任：\\n与阿努拉及北语一女生代表北京队参加北京电视台《第三届中国通电视大赛》（2000年9月）\\n江苏卫视的《青春大碰撞》作嘉宾主持（2001年9月-12月）\\n代表江苏队参加《春节外国人中华才艺大赛》的戏曲和曲艺项目，获得最佳表演奖（2002年9月）\\n中央电视台国际频道《同乐五洲》主持（2002年8月）\\n中央电视台国际频道《学汉语─快乐中国》主持（2003年3月）",
  "教育信息": "切尔滕纳姆伯恩赛德学校（1992年9月～1997年7月）\\n英国切尔滕纳姆佩茨文法学校六年级（高中部）（1997年9月～1999年7月）\\n在杜伦大学东亚研究系学习汉语（1999年10月～2000年7月）\\n中国人民大学进行一年的汉语培训（2000年9月～2001年7月）\\n中国人民大学 文学学士（2000年 - 2004年）\\n中国人民大学 法学院 — 法学硕士 LLM（2014年 - 2016年）\\n清华大学博士研究生(2020年9月 - 至今)",
  "社交网络": "与中国人民大学保持紧密联系，作为校友代表出席国际文化交流学院成立大会并发表讲话。",
  "社会影响": "被中国媒体称为‘中国通’，参与媒体与活动传播中华文化；被誉为“最会说汉语的外国主持人”。",
  "个人重要成就/著作": "主持中央电视台国际频道节目《同乐五洲》《快乐中国——学汉语》；获得中国人民大学文学学士与法学硕士学位；2025年参与纪录片《文运中国》拍摄，推动中国文化国际传播。",
  "关键事件": "2000–2001年：赴中国人民大学学习中文语言文学，开启在华生活。\\n2004年起：加入中央电视台国际频道，主持《同乐五洲》。\\n2014年：进入中国人民大学法学院攻读法学硕士。\\n2025年：与法国主持人克莱奥·吕登共同出演纪录片《文运中国》。",
  "重要活动": "2014 年倡导汉语学习与文化理解；2025 年通过纪录片《文运中国》呼吁全球观众亲身感受中国。"
}
""",
        "深度画像": """
{
  "兴趣偏好": "表演话剧，演电影弹钢琴，吹小号，弹古琴，喜爱戏曲",
  "性格类型倾向": "外向开放型",
  "性格弱点": "具有明显的理想主义倾向、舆论压力大、自我要求高",
  "价值取向": "文化认同：认为自己是一个“中国主义者”，推崇中国文化与传统艺术（书法、国学、戏曲等）",
  "职业路径分析": "文化传播型：语言文化深耕（1999年起）→ 传媒传播实践（2003年起，进入央视）→ 文化交流阶段（至今，活跃于高校与国际交流活动）。",
  "未来发展预期": "友华人物：预计未来也会与中国友好；文化传播者：希望未来能“邀请更多全球观众来中国”，继续强调跨文化理解。"
}
"""
    }

# ================= 辅助组件类 =================
class KeywordExtractor:
    @staticmethod
    def extract(text: str, top_k: int = 10) -> List[str]:
        if not text: return []
        # TextRank 适合提取名词短语，更适合作为关键词
        keywords = jieba.analyse.textrank(text, topK=top_k, withWeight=False, allowPOS=('ns', 'n', 'vn', 'nr', 'nt'))
        if len(keywords) < 2:
            keywords = jieba.analyse.extract_tags(text, topK=top_k)
        return keywords

class ReciprocalRankFusion:
    def __init__(self, k=60): self.k = k
    
    def fuse(self, ranked_lists, weights=None):
        """支持加权的 RRF 融合"""
        if not ranked_lists: return []
        
        if weights is None:
            weights = [1.0] * len(ranked_lists)
            
        if len(weights) != len(ranked_lists):
            weights = [1.0] * len(ranked_lists)

        scores = {}
        all_nodes = {}
        
        for lst, w in zip(ranked_lists, weights):
            for rank, node in enumerate(lst):
                nid = node.node.node_id
                all_nodes[nid] = node.node
                # 加权 RRF 公式
                scores[nid] = scores.get(nid, 0) + w * (1.0 / (self.k + rank + 1))
        
        fused = [NodeWithScore(node=all_nodes[nid], score=score) for nid, score in scores.items()]
        fused.sort(key=lambda x: x.score, reverse=True)
        return fused

# ================= 模型封装类 (Local) =================
class Qwen3Embedding:
    """本地 Embedding 模型"""
    def __init__(self, model_name):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
        self.model = AutoModel.from_pretrained(model_name, trust_remote_code=True, torch_dtype=torch.float16, device_map=EnhancedConfig.DEVICE)
        self.model.eval()

    def get_text_embeddings(self, texts):
        with torch.no_grad():
            inputs = self.tokenizer(texts, padding=True, truncation=True, return_tensors="pt", max_length=512).to(EnhancedConfig.DEVICE)
            outputs = self.model(**inputs)
            embeddings = self._mean_pooling(outputs.last_hidden_state, inputs['attention_mask'])
        return embeddings.cpu().numpy().tolist()
    
    def get_query_embedding(self, query): return self.get_text_embeddings([query])[0]
    
    def _mean_pooling(self, model_output, attention_mask):
        token_embeddings = model_output
        input_mask_expanded = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
        return torch.sum(token_embeddings * input_mask_expanded, 1) / torch.clamp(input_mask_expanded.sum(1), min=1e-9)

class Qwen3DirectReranker:
    """本地 Reranker 模型"""
    def __init__(self, model_name):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
        self.model = AutoModelForCausalLM.from_pretrained(model_name, trust_remote_code=True, torch_dtype=torch.float16, device_map=EnhancedConfig.DEVICE)
        self.model.eval()

    def rerank(self, query, nodes, top_n):
        if not nodes: return []
        scored_nodes = []
        for node in nodes:
            text_snippet = node.node.text[:512]
            prompt = f"查询: {query}\n文档: {text_snippet}"
            inputs = self.tokenizer(prompt, return_tensors="pt").to(EnhancedConfig.DEVICE)
            with torch.no_grad():
                outputs = self.model(**inputs)
                logits = outputs.logits[:, -1, :]
                yes_id = self.tokenizer.encode("yes", add_special_tokens=False)[0]
                no_id = self.tokenizer.encode("no", add_special_tokens=False)[0]
                score = np.exp(logits[0, yes_id].item()) / (np.exp(logits[0, yes_id].item()) + np.exp(logits[0, no_id].item()) + 1e-9)
            node.score = score
            scored_nodes.append(node)
        return sorted(scored_nodes, key=lambda x: x.score, reverse=True)[:top_n]

class SimpleLLM:
    """本地 vLLM 模型"""
    def __init__(self, model_path):
        self.tokenizer = AutoTokenizer.from_pretrained(model_path, trust_remote_code=True)
        self.llm = LLM(model=model_path, 
                       gpu_memory_utilization=EnhancedConfig.gpu_memory_utilization, 
                       max_model_len=EnhancedConfig.max_model_len, 
                       trust_remote_code=True,
                       tensor_parallel_size=1)
        self.sampling_params = SamplingParams(
            temperature=EnhancedConfig.TEMPERATURE, 
            max_tokens=EnhancedConfig.MAX_NEW_TOKENS, 
            top_p=EnhancedConfig.TOP_P
        )

    def chat(self, messages: List[ChatMessage]) -> ChatMessage:
        prompt = self.tokenizer.apply_chat_template(
            [{"role": m.role, "content": m.content} for m in messages], 
            tokenize=False, 
            add_generation_prompt=True
        )
        outputs = self.llm.generate([prompt], self.sampling_params)
        return ChatMessage(role="assistant", content=outputs[0].outputs[0].text)

# ================= 模型封装类 (API) =================

class OpenAIEmbedding:
    """API 模式: Embedding"""
    def __init__(self, api_key, base_url, model_name):
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.model_name = model_name

    def get_text_embeddings(self, texts: List[str]) -> List[List[float]]:
        # OpenAI API 有 batch size 限制，通常一次不超过 2048 个 token 或一定数量的 string
        # 这里简单处理，如果 texts 太多可以自行分批
        texts = [t.replace("\n", " ") for t in texts] # 推荐操作：移除换行
        response = self.client.embeddings.create(input=texts, model=self.model_name)
        return [data.embedding for data in response.data]

    def get_query_embedding(self, query: str) -> List[float]:
        return self.get_text_embeddings([query])[0]

class OpenAILLM:
    """API 模式: LLM"""
    def __init__(self, api_key, base_url, model_name):
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.model_name = model_name

    def chat(self, messages: List[ChatMessage]) -> ChatMessage:
        # 将 LlamaIndex 的 ChatMessage 转换为 OpenAI 的 dict 格式
        openai_messages = [{"role": m.role, "content": m.content} for m in messages]
        
        response = self.client.chat.completions.create(
            model=self.model_name,
            messages=openai_messages,
            temperature=EnhancedConfig.TEMPERATURE,
            max_tokens=EnhancedConfig.MAX_NEW_TOKENS,
            top_p=EnhancedConfig.TOP_P
        )
        content = response.choices[0].message.content
        return ChatMessage(role="assistant", content=content)

class APIReranker:
    """API 模式: Reranker (通用 HTTP 接口)"""
    def __init__(self, api_key, api_url, model_name):
        self.api_key = api_key
        self.api_url = api_url
        self.model_name = model_name

    def rerank(self, query, nodes, top_n):
        if not nodes: return []
        
        # 提取文档文本列表
        documents = [n.node.text for n in nodes]
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": self.model_name,
            "query": query,
            "documents": documents,
            "top_n": top_n
        }
        
        try:
            response = requests.post(self.api_url, json=payload, headers=headers, timeout=10)
            response.raise_for_status()
            results = response.json().get("results", [])
            
            # 将 API 返回结果映射回 NodeWithScore
            scored_nodes = []
            for res in results:
                idx = res["index"]
                score = res["relevance_score"]
                node = nodes[idx]
                node.score = score
                scored_nodes.append(node)
            
            return sorted(scored_nodes, key=lambda x: x.score, reverse=True)
            
        except Exception as e:
            print(f"⚠️ Rerank API 调用失败: {e}，将返回原始顺序的前 {top_n} 个。")
            return nodes[:top_n]

# ================= 数据处理 =================
class GlobalDocumentStore:
    """
    全局文档存储
    _store 结构优化为: { doc_id: {'text': str, 'metadata': dict} }
    以支持返回完整 Node 对象给 _build_context_str
    """
    _store = {} 
    
    @classmethod
    def add_document(cls, doc_id, text, metadata=None):
        if metadata is None: metadata = {}
        cls._store[doc_id] = {'text': text, 'metadata': metadata}
        
    @classmethod
    def get_document_data(cls, doc_id):
        """返回包含文本和元数据的完整数据包"""
        return cls._store.get(doc_id, {'text': "", 'metadata': {}})
    
    @classmethod
    def get_document(cls, doc_id):
        """兼容旧接口，只返回文本"""
        return cls._store.get(doc_id, {}).get('text', "")

class EnhancedDocumentProcessor:
    def read_all_documents(self, base_path):
        all_docs = []
        data_path = os.path.join(base_path, "清洗数据")
        if not os.path.exists(data_path):
            data_path = os.path.join(base_path, "data")
            
        print(f"正在从 {data_path} 读取文档...")
        for root, _, files in os.walk(data_path):
            for file in files:
                fpath = os.path.join(root, file)
                try:
                    content = ""
                    if file.endswith('.txt') or file.endswith('.md') or file.endswith('.csv'):
                        with open(fpath, 'r', errors='ignore', encoding='utf-8') as f: content = f.read()
                    elif file.endswith('.docx'):
                        doc = docx.Document(fpath)
                        content = "\n".join([p.text for p in doc.paragraphs])
                    
                    if content.strip():
                        doc_id = str(uuid.uuid4())
                        # 确保元数据中包含 file_name
                        metadata = {"doc_id": doc_id, "file_name": file}
                        doc_obj = Document(text=content, metadata=metadata)
                        all_docs.append(doc_obj)
                        
                        # 存入全局 Store
                        GlobalDocumentStore.add_document(doc_id, content, metadata=metadata)
                        
                except Exception as e:
                    print(f"Skipped {file}: {e}")
        return all_docs

class SimpleVectorStore:
    def __init__(self, nodes, embed_model, index_path):
        self.nodes = nodes
        self.embed_model = embed_model
        self.index_path = index_path
        if os.path.exists(index_path):
            print(f"加载现有向量索引: {index_path}")
            self.index = faiss.read_index(index_path)
        else:
            self._build_index()

    def _build_index(self):
        print("构建向量索引...")
        if not self.nodes: return
        batch_size = 32
        texts = [n.text for n in self.nodes]
        embeddings = []
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i+batch_size]
            embeddings.extend(self.embed_model.get_text_embeddings(batch))
        arr = np.array(embeddings, dtype=np.float32)
        faiss.normalize_L2(arr)
        self.index = faiss.IndexFlatIP(arr.shape[1])
        self.index.add(arr)
        os.makedirs(os.path.dirname(self.index_path), exist_ok=True)
        faiss.write_index(self.index, self.index_path)

    def search(self, query, top_k):
        if not self.nodes: return []
        q_emb = np.array([self.embed_model.get_query_embedding(query)], dtype=np.float32)
        faiss.normalize_L2(q_emb)
        sims, idxs = self.index.search(q_emb, top_k)
        return [NodeWithScore(node=self.nodes[i], score=float(s)) for s, i in zip(sims[0], idxs[0]) if i != -1]

# ================= 核心系统类 =================
class EnhancedHybridRetriever:
    def __init__(self, vector_store, nodes, reranker):
        self.vector_store = vector_store
        self.nodes = nodes
        self.reranker = reranker
        self.bm25 = None
        if EnhancedConfig.BM25_ENABLED and nodes:
            print("构建 BM25 索引...")
            tokenized_corpus = [jieba.lcut(n.text) for n in nodes]
            self.bm25 = BM25Okapi(tokenized_corpus)
        self.rrf = ReciprocalRankFusion(k=EnhancedConfig.RRF_K)

    def retrieve(self, query_name, original_query):
        # 1. 统一转为列表处理
        queries = original_query if isinstance(original_query, list) else [original_query]
        rerank_query_text = queries[0] 

        # 定义融合权重
        VECTOR_WEIGHT = 1.0
        BM25_WEIGHT = 0.5

        # 阶段 1: 多路混合检索
        all_result_lists = []
        all_weights = []

        for q in queries:
            # A. 向量检索
            vec_nodes = self.vector_store.search(q, EnhancedConfig.SIMILARITY_TOP_K)
            all_result_lists.append(vec_nodes)
            all_weights.append(VECTOR_WEIGHT)
            
            # B. BM25 检索
            if self.bm25:
                if any('\u4e00' <= char <= '\u9fff' for char in q):
                    tokenized_query = jieba.lcut(q)
                else:
                    tokenized_query = q.split() 
                
                scores = self.bm25.get_scores(tokenized_query)
                top_n_indices = np.argsort(scores)[::-1][:EnhancedConfig.BM25_TOP_K]
                
                bm25_nodes = [NodeWithScore(node=self.nodes[i], score=float(scores[i])) 
                              for i in top_n_indices if scores[i] > 1.0]
                
                all_result_lists.append(bm25_nodes)
                all_weights.append(BM25_WEIGHT)
        
        # C. 加权 RRF 融合
        chunk_candidates_list = self.rrf.fuse(all_result_lists, weights=all_weights)
        
        # 阶段 2: 关键词元数据增强
        final_scored_chunks = []
        for item in chunk_candidates_list:
            node = item.node
            score = item.score
            node_keywords = node.metadata.get('keywords', [])
            
            hit_keyword = False
            for kw in node_keywords:
                if query_name in kw or kw in query_name:
                    hit_keyword = True
                    break
            
            if hit_keyword:
                score *= 2.0 
            
            item.score = score
            final_scored_chunks.append(item)
            
        final_scored_chunks.sort(key=lambda x: x.score, reverse=True)

        # 阶段 3: 重排序
        if self.reranker:
            final_scored_chunks = self.reranker.rerank(rerank_query_text, final_scored_chunks, top_n=20)

        # 阶段 4: 父文档映射 (Small-to-Big)
        doc_scores = {}
        for item in final_scored_chunks:
            doc_id = item.node.metadata.get('doc_id')
            if not doc_id: continue
            if item.score > doc_scores.get(doc_id, 0):
                doc_scores[doc_id] = item.score

        sorted_doc_ids = sorted(doc_scores.items(), key=lambda x: x[1], reverse=True)[:EnhancedConfig.RERANK_TOP_N]
        
        # 【关键修改】返回 NodeWithScore 对象列表，以便 _build_context_str 处理
        final_parent_nodes = []
        for doc_id, score in sorted_doc_ids:
            doc_data = GlobalDocumentStore.get_document_data(doc_id)
            if not doc_data.get('text'): continue
            
            # 重建 Node (带元数据)
            full_node = Document(text=doc_data['text'], metadata=doc_data.get('metadata', {}))
            final_parent_nodes.append(NodeWithScore(node=full_node, score=score))
            
        return final_parent_nodes

class FullyOptimizedQASystem:
    def __init__(self):
        self.vector_store = None
        self.retriever = None
        self.llm = None
        self.embed_model = None
        self.reranker = None
        self.nodes = []
        self.faiss_path = '../data/FAISS_Vector_DB/faiss_index.bin'

    def initialize(self, doc_path: str):
        print(f"1. 初始化系统 (模式: {EnhancedConfig.MODE})...")
        
        # === 1. 加载模型 (Local vs API) ===
        if EnhancedConfig.MODE == "local":
            print("   正在加载本地模型 (Qwen3)...")
            self.embed_model = Qwen3Embedding(EnhancedConfig.EMBEDDING_MODEL)
            self.llm = SimpleLLM(EnhancedConfig.GENERATION_MODEL)
            if EnhancedConfig.RERANKER_ENABLED:
                self.reranker = Qwen3DirectReranker(EnhancedConfig.RERANKER_MODEL)
        else:
            print("   正在初始化 API 客户端...")
            if OpenAI is None: raise ImportError("使用 API 模式请先安装 `pip install openai`")
            
            self.embed_model = OpenAIEmbedding(
                api_key=EnhancedConfig.API_EMBED_KEY,
                base_url=EnhancedConfig.API_EMBED_BASE_URL,
                model_name=EnhancedConfig.API_EMBED_MODEL_NAME
            )
            self.llm = OpenAILLM(
                api_key=EnhancedConfig.API_KEY,
                base_url=EnhancedConfig.API_BASE_URL,
                model_name=EnhancedConfig.API_LLM_MODEL_NAME
            )
            if EnhancedConfig.RERANKER_ENABLED:
                self.reranker = APIReranker(
                    api_key=EnhancedConfig.API_RERANK_KEY,
                    api_url=EnhancedConfig.API_RERANK_URL,
                    model_name=EnhancedConfig.API_RERANK_MODEL_NAME
                )

        # === 2. 缓存处理与文档加载 ===
        # --- 缓存逻辑开始 ---
        os.makedirs(EnhancedConfig.CACHE_DIR, exist_ok=True)
        nodes_cache_path = os.path.join(EnhancedConfig.CACHE_DIR, EnhancedConfig.NODES_CACHE_FILE)
        doc_store_cache_path = os.path.join(EnhancedConfig.CACHE_DIR, EnhancedConfig.DOC_STORE_CACHE_FILE)
        
        # 尝试加载缓存
        if os.path.exists(nodes_cache_path) and os.path.exists(doc_store_cache_path):
            print(f"2. 检测到节点缓存，正在加载: {nodes_cache_path} ...")
            try:
                with open(nodes_cache_path, 'rb') as f:
                    self.nodes = pickle.load(f)
                with open(doc_store_cache_path, 'rb') as f:
                    GlobalDocumentStore._store = pickle.load(f)
                print(f"✓ 成功加载 {len(self.nodes)} 个节点。")
            except Exception as e:
                print(f"⚠️ 缓存加载失败 ({e})，将重新处理文档...")
                self._process_and_cache_docs(doc_path, nodes_cache_path, doc_store_cache_path)
        else:
            print("2. 未检测到缓存，开始全量处理文档...")
            self._process_and_cache_docs(doc_path, nodes_cache_path, doc_store_cache_path)
        # --- 缓存逻辑结束 ---

        print("4. 构建/加载向量索引 (FAISS)...")
        self.vector_store = SimpleVectorStore(self.nodes, self.embed_model, self.faiss_path)
        
        self.retriever = EnhancedHybridRetriever(self.vector_store, self.nodes, self.reranker)
        print("✓ 系统初始化完成")

    def _process_and_cache_docs(self, doc_path, nodes_path, doc_store_path):
        """处理文档、提取关键词并保存缓存的辅助函数"""
        print("   (a) 读取源文件...")
        processor = EnhancedDocumentProcessor()
        docs = processor.read_all_documents(doc_path)
        
        print("   (b) 文档切分与关键词提取 (此步骤较慢)...")
        splitter = SentenceSplitter(chunk_size=EnhancedConfig.CHUNK_SIZE, chunk_overlap=EnhancedConfig.CHUNK_OVERLAP)
        self.nodes = []
        
        for d in docs:
            doc_id = d.metadata['doc_id']
            cur_nodes = splitter.get_nodes_from_documents([d])
            for n in cur_nodes:
                n.metadata['doc_id'] = doc_id
                # 确保传递 file_name，供 _build_context_str 使用
                n.metadata['file_name'] = d.metadata.get('file_name', 'unknown')
                n.metadata['keywords'] = KeywordExtractor.extract(n.text, top_k=5)
            self.nodes.extend(cur_nodes)
            
        print(f"   (c) 生成 {len(self.nodes)} 个索引节点。")
        
        print("   (d) 正在保存缓存到磁盘...")
        with open(nodes_path, 'wb') as f:
            pickle.dump(self.nodes, f)
        with open(doc_store_path, 'wb') as f:
            pickle.dump(GlobalDocumentStore._store, f)
        print("✓ 缓存保存完成。")

    # --- 以下为核心业务逻辑 ---

    def _post_process(self, answer: str) -> str:
        pattern = r'<think>(.*?)</think>'
        content = re.sub(pattern, '', answer, flags=re.DOTALL)
        return re.sub(r'\n+', '\n', content).strip()

    def _safe_llm(self, prompt, label):
        attempt = 0
        while True:
            attempt += 1
            try:
                raw_res = self.llm.chat([ChatMessage(role="user", content=prompt)]).content
                processed_res = self._post_process(raw_res)
                json_str = self._clean_json(processed_res)
                return json.loads(json_str)
            except Exception as e:
                print(f"❌ {label} 解析失败 (第 {attempt} 次尝试): {e}")
                if attempt >= 5: 
                     print("🚫 达到最大重试次数，跳过。")
                     return {}
                continue

    def _clean_json(self, text):
        text = re.sub(r'```json\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'```', '', text)
        s = text.find('{')
        e = text.rfind('}')
        if s != -1 and e != -1:
            return text[s:e+1]
        return "{}"

    def extract_info(self, name: str, context: str) -> Dict:
        all_res = {}
        field_definitions = {
            '姓名（原文）': f"人物【{name}】的外文全名、原名或曾用名。非中文名必须提取。",
            '出生年份(如1950年)': "仅提取数字年份（例如：1965）。",
            '家庭关系（如“关系：人名”）': "配偶、子女、父母等。格式：'关系：姓名'。",
            '社会关系（如“关系：人名”）': "恩师、密友、合作伙伴。格式：'关系：姓名'。",
            '任职信息': "曾任及现任职务。格式：'时间：机构/职位'。多项用分号分隔。",
            '教育信息': "学位、毕业院校及专业。包含留学经历。",
            '社交网络': "所属圈层、俱乐部或核心人脉网络。",
            '社会影响': "影响力评估、荣誉头衔。",
            '个人重要成就/著作': "代表性作品、成果或奖项。",
            '关键事件': "转折性事件。",
            '兴趣偏好': "业余爱好、生活习惯。",
            '性格类型倾向': "根据言行推断的性格特征。",
            '性格弱点': "根据争议事件推断的性格短板。",
            '价值取向': "公开表达或行动体现的核心价值观。",
            '职业路径分析': "职业上升逻辑总结。",
            '未来发展预期': "基于现状对未来的预测。"
        }

        for group_name, fields in EnhancedConfig.EXTRACT_GROUPS.items():
            current_field_descs = {k: field_definitions.get(k, "根据上下文提取") for k in fields}
            example_json = EnhancedConfig.FEW_SHOT_EXAMPLES.get(group_name, "{}")
            
            group_instruction = ""
            if group_name == "基础身份":
                group_instruction = "【指令】：注重准确性，严格区分同名人物。"
            elif group_name == "生涯成就":
                group_instruction = "【指令】：注重时间线（倒序或顺序），多项内容用分号分隔。"
            elif group_name == "深度画像":
                group_instruction = "【指令】：需要基于行为进行推理（Inference），不要仅摘抄。"

            prompt = f"""你是一位高级情报分析师。基于资料构建人物【{name}】的【{group_name}】档案。

【背景资料】：
{context[:EnhancedConfig.MAX_CONTEXT_LENGTH]}

【待提取字段及定义】：
{json.dumps(current_field_descs, ensure_ascii=False, indent=2)}

【参考范例 (Style Guide)】：
请严格模仿以下 JSON 的字段填写风格（尤其是时间线格式和分号分隔）：
{example_json}

{group_instruction}

【要求】：
1. 仅输出 JSON。
2. 缺失字段填“未知”。
3. 多条信息用分号分隔。

请生成 JSON：
"""
            res = self._safe_llm(prompt, group_name)
            for k in fields:
                if k not in res: res[k] = "未知"
            all_res.update(res)
            
        all_res['姓名（中文）'] = name
        return all_res

    def _translate_name(self, name: str) -> str:
        """
        优化：智能判断语种的翻译 Prompt
        根据人物背景决定是输出俄语还是英语
        """
        prompt = f"""Task: Identify the real-world person associated with the Chinese name '{name}'.

Instructions:
1. If the person is from a **Russian-speaking country** (e.g., Russia, Ukraine, Belarus, USSR), output their name in **Russian (Cyrillic)**.
2. For all other persons (Western, International, etc.), output their name in **English**.
3. Output **ONLY** the name. Do not include any explanation, punctuation, or extra words.

Target Name:"""
        try:
            raw_res = self.llm.chat([ChatMessage(role="user", content=prompt)]).content
            process_res = self._post_process(raw_res)
            # 简单的清洗，移除句号
            return process_res.strip().replace(".", "").replace("。", "")
        except:
            return name
    
    # 【新增】构建上下文的模块，控制长度并保留来源
    def _build_context_str(self, nodes):
        context = []
        added_ids = set()
        cur_len = 0
        
        # 传入的 nodes 已经是 NodeWithScore 列表，按分数排序
        sorted_nodes = sorted(nodes, key=lambda x: x.score, reverse=True)
        
        for node_score in sorted_nodes:
            node = node_score.node
            # 使用 file_name 和 doc_id 联合去重
            key = f"{node.metadata.get('file_name', 'unknown')}_{node.node_id}"
            if key in added_ids: continue
            
            # 格式化文本
            text = f"【来源:{node.metadata.get('file_name', '未知')}】\n{node.text}\n"
            
            # 严格控制长度
            if cur_len + len(text) > EnhancedConfig.MAX_CONTEXT_LENGTH: break
            
            context.append(text)
            added_ids.add(key)
            cur_len += len(text)
            
        return "\n".join(context)

    def generate_person_profile(self, name: str) -> Dict:
        print(f"\n🚀 开始生成人物画像: {name}")
        
        # 1. 翻译名字 (俄语/英语 智能判断)
        translated_name = self._translate_name(name)
        print(f"🔤 翻译结果: {translated_name}")

        # 2. 构建多维度查询列表 (恢复了被注释的内容)
        queries = [
            f"{name}",
            f"{translated_name}"
        ]
        
        print(f"🔍 执行多路检索, Query 数量: {len(queries)}")
        
        # 3. 检索 (返回 NodeWithScore 列表)
        retrieved_nodes = self.retriever.retrieve(query_name=name, original_query=queries)
        
        # 4. 构建上下文 (使用 _build_context_str)
        context_str = self._build_context_str(retrieved_nodes)
        
        print(f"📄 最终构建上下文长度: {len(context_str)} 字符")
        if len(context_str) < 50:
            print("⚠️ 警告: 未检索到有效内容！")
        
        # 5. 提取信息
        profile_data = self.extract_info(name, context_str)
        return profile_data