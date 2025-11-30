import os
import re
import base64
import fitz  # PyMuPDF
import csv
import pickle
from docx import Document
from openai import OpenAI

# 引入配置
from settings import CONF

# ================= 初始化 =================
client = OpenAI(api_key=CONF['api']['key'], base_url=CONF['api']['base_url'])
DEVICE = CONF['system']['device']
print(f"当前运行设备: {DEVICE.upper()}")

# ---------------------------------------------------------
# 模块 1: 综合清洗管道 (已更新)
# ---------------------------------------------------------

def clean_ocr_artifacts(text):
    """
    专门针对 Layout/OCR 模型输出的特殊标记进行清洗。
    移除 <|ref|>...<|/ref|>, <|det|>...<|/det|> 以及 HTML 标签，只保留内容。
    """
    if not text: return ""

    # 1. 移除成对出现的 OCR 元数据块及其内容 (如 <|ref|>title<|/ref|>)
    #    这些通常是描述标签，而不是正文内容
    text = re.sub(r'<\|(ref|det|image_caption_title)\|>(.*?)<\|\/\1\|>', '', text, flags=re.DOTALL)

    # 2. 移除独立的 OCR 特殊 Token (如 <|image|>, <|box|>)
    text = re.sub(r'<\|.*?\|>', '', text)

    # 3. 移除 Bounding Box 坐标 (如 [[57, 87, 460, 120]])
    text = re.sub(r'\[\[\d+,\s*\d+,\s*\d+,\s*\d+\]\]', '', text)

    # 4. 移除 HTML 标签但保留内容 (如 <center>...</center> -> ...)
    text = re.sub(r'<[^>]+>', '', text)

    return text

def is_citation_heavy(text, threshold=None):
    """判断是否为高密度参考文献页"""
    if threshold is None:
        threshold = CONF['rag']['citation_threshold']
        
    if not text: return False
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines: return False
    
    citation_score = 0
    total_lines = len(lines)
    
    keywords = [r"doi\.?", r"http", r"vol\.", r"pp\.", r"et\s+al", r"isbn", r"issn", r"arxiv"]
    keyword_pattern = re.compile("|".join(keywords), re.IGNORECASE)
    year_pattern = re.compile(r"[\(\[\s](19|20)\d{2}[\)\]\.]")
    ref_index_pattern = re.compile(r"^\[\d+\]|^\d+\.\s+[A-Z]")

    for line in lines:
        if year_pattern.search(line) or keyword_pattern.search(line) or (ref_index_pattern.match(line) and len(line) > 20):
            citation_score += 1
            
    return (citation_score / total_lines) > threshold

def clean_web_noise(text):
    """针对百度百科、网页复制内容的特定清洗"""
    if not text: return ""
    lines = text.split('\n')
    cleaned_lines = []
    
    ignore_keywords = [
        "登录", "注册", "秒懂百科", "特色百科", "加入百科", "上传视频", 
        "同名词条", "查看更多", "相关星图", "词条统计", "浏览次数", 
        "编辑次数", "最近更新", "播报", "回忆格式", "ppt模板", "下载谷歌"
    ]
    
    sidebar_pattern = re.compile(r'^\d+\s*[:\s]\s*[\u4e00-\u9fa5a-zA-Z0-9]+$')
    
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        
        # 允许短标题（如Markdown标题），但过滤过短的无意义字符
        if len(stripped) < 2 and not stripped.startswith("#"):
            continue
            
        if any(kw in stripped for kw in ignore_keywords):
            continue
            
        if sidebar_pattern.match(stripped):
            continue
        
        stripped = stripped.replace("展开40个", "").replace("收起", "")
        cleaned_lines.append(stripped)
        
    return "\n".join(cleaned_lines)

def advanced_clean_pipeline(text):
    """综合清洗入口"""
    if not text: return ""

    # Step 0: 先移除 OCR 产生的特殊字符和元数据
    text = clean_ocr_artifacts(text)

    # Step 1: 检查是否为参考文献页
    if is_citation_heavy(text):
        return "<<CITATION_PAGE>>"
    
    # Step 2: 清洗网页/通用噪声
    text = clean_web_noise(text)

    # Step 3: 格式规范化
    lines = text.split('\n')
    final_lines = []
    for line in lines:
        s = line.strip()
        # 移除纯页码
        if re.match(r'^[-—\s]*\d+[-—\s]*$', s) or re.match(r'^Page\s+\d+', s, re.IGNORECASE):
            continue
        if s:
            final_lines.append(s)
            
    return "\n\n".join(final_lines)

# ---------------------------------------------------------
# 模块 2: 多格式文档处理
# ---------------------------------------------------------

def image_to_base64(image_bytes):
    return base64.b64encode(image_bytes).decode("utf-8")

def process_pdf_page_ocr(image_bytes):
    try:
        response = client.chat.completions.create(
            model=CONF['api']['ocr_model'],
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": "data:image/png;base64," + image_to_base64(image_bytes)}},
                    {"type": "text", "text": "<image>\n<|grounding|>Convert the document to markdown. "}
                ]
            }]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"  [OCR Error]: {e}")
        return ""

def convert_pdf_to_md(filepath):
    full_text = ""
    doc = fitz.open(filepath)
    print(f"处理 PDF: {os.path.basename(filepath)} (共 {doc.page_count} 页)")
    
    found_ref = False
    for page_num in range(doc.page_count):
        if found_ref: break
        
        print(f"  > OCR 第 {page_num+1} 页...")
        pix = doc.load_page(page_num).get_pixmap(matrix=fitz.Matrix(2, 2))
        raw = process_pdf_page_ocr(pix.tobytes("png"))
        
        # 将原始 OCR 结果送入清洗管道
        cleaned = advanced_clean_pipeline(raw)
        
        if cleaned == "<<CITATION_PAGE>>":
            print("  [熔断] 发现参考文献页，停止处理该文件后续页面。")
            found_ref = True
            continue
            
        full_text += cleaned + "\n\n"
    return full_text

def convert_docx_to_md(filepath):
    print(f"处理 DOCX: {os.path.basename(filepath)}")
    try:
        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text: full_text.append(" | ".join(row_text))
        
        raw = "\n".join(full_text)
        cleaned = advanced_clean_pipeline(raw)
        return "" if cleaned == "<<CITATION_PAGE>>" else cleaned
    except Exception as e:
        print(f"  [DOCX Error] {e}")
        return ""

def convert_csv_to_md(filepath):
    print(f"处理 CSV: {os.path.basename(filepath)}")
    try:
        lines = []
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            headers = next(reader, None)
            if headers:
                lines.append(" | ".join(headers))
                lines.append("|".join(["---"] * len(headers)))
            
            for row in reader:
                clean_row = [str(item).strip().replace('\n', ' ') for item in row]
                lines.append(" | ".join(clean_row))
        
        return "\n".join(lines)
    except Exception as e:
        print(f"  [CSV Error] {e}")
        return ""

def process_documents():
    print("=== 步骤 1/2: 文档转换与清洗 ===")
    input_folder = CONF['paths']['input_folder']
    processed_folder = CONF['paths']['processed_folder']

    if not os.path.exists(input_folder): return False
    if not os.path.exists(processed_folder): os.makedirs(processed_folder)

    for filename in os.listdir(input_folder):
        input_path = os.path.join(input_folder, filename)
        output_path = os.path.join(processed_folder, os.path.splitext(filename)[0] + ".md")
        
        if os.path.exists(output_path):
            print(f"跳过已存在: {filename}")
            continue

        content = ""
        ext = filename.lower()
        if ext.endswith(".pdf"): content = convert_pdf_to_md(input_path)
        elif ext.endswith(".docx"): content = convert_docx_to_md(input_path)
        elif ext.endswith(".csv"): content = convert_csv_to_md(input_path)
            
        if content:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
                
if __name__ == "__main__":
    process_documents()