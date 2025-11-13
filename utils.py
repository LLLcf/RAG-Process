import os
import re
import json
import time
import math
import asyncio
import warnings
from typing import List, Dict, Tuple, Optional, AsyncIterator, Iterator, AsyncGenerator, Set, Any
from pathlib import Path
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, as_completed

import pandas as pd
import numpy as np
from tqdm import tqdm
import torch
import torch.nn.functional as F
from rank_bm25 import BM25Okapi
import jieba
import docx
from transformers import AutoModel, AutoTokenizer, AutoModelForCausalLM

# LlamaIndex核心组件
from llama_index.core import (
    VectorStoreIndex,
    Document,
    Settings,
    StorageContext,
    ServiceContext
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.postprocessor import SimilarityPostprocessor
from llama_index.core.schema import NodeWithScore, QueryBundle, BaseNode
from llama_index.core.llms import (
    LLM, ChatMessage, CompletionResponse, 
    CompletionResponseGen, ChatResponse, ChatResponseGen
)

import os
import numpy as np
import faiss
from tqdm import tqdm
from typing import List, Optional
import jieba
from rank_bm25 import BM25Okapi
import hashlib
from openai import OpenAI
from vllm import LLM, SamplingParams

# 忽略警告
warnings.filterwarnings('ignore')

@dataclass
class EnhancedConfig:
    """增强配置类"""
    
    # 模型路径
    EMBEDDING_MODEL = "/root/lanyun-fs/models/Qwen3-Embedding-0.6B"
    RERANKER_MODEL = "/root/lanyun-fs/models/Qwen3-Reranker-0.6B"
    # GENERATION_MODEL = "/root/lanyun-fs/models/Qwen3-0.6B"
    GENERATION_MODEL = "/root/lanyun-tmp/models/Qwen3-4B"
    
    # 指令
    EMBEDDING_INSTRUCTION = ""
    RERANKER_INSTRUCTION = ""
    
    # 分割参数 - 改为段落分割
    PARAGRAPH_SEPARATOR = "\n\n"
    MAX_PARAGRAPH_LENGTH = 1024
    MIN_PARAGRAPH_LENGTH = 200
    
    # 分块参数
    CHUNK_SIZE = 1024
    CHUNK_OVERLAP = 128
    MIN_CHUNK_SIZE = 100
    
    # 分割模式选择：'paragraph' 或 'sentence'
    CHUNK_MODE = 'sentence'
    
    # 检索参数
    SIMILARITY_TOP_K = 50
    BM25_TOP_K = 50
    RERANK_TOP_N = 25
    FINAL_TOP_K = 5

    include_knowledge_graph = True
    
    # 组件开关
    QUERY_REWRITE_ENABLED = True
    QUERY_DECOMPOSE_ENABLED = False
    HYPO_ANSWER_ENABLED = False
    DEDUPLICATE_ENABLED = True
    RERANKER_ENABLED = True
    BM25_ENABLED = True
    
    # 多查询参数
    QUERY_REWRITE_NUM = 3
    
    # RRF参数
    RRF_K = 60
    RRF_ENABLED = True
    # 上下文窗口
    CONTEXT_WINDOW = 1
    
    # 生成参数
    MAX_NEW_TOKENS = 5000
    TEMPERATURE = 0.3
    TOP_P = 1.0
    gpu_memory_utilization = 0.9
    max_model_len = 36000
    max_context_length = 30000
    # 设备配置
    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
    
    # SimHash去重参数
    HASH_BITS = 64
    SIMILAR_THRESHOLD = 5
    MIN_TEXT_LENGTH = 100

class KnowledgeGraphProcessor:
    """知识图谱数据处理器"""
    
    def __init__(self):
        self.documents = []
    
    def load_excel_knowledge_graph(self, excel_path: str) -> List[Document]:
        """加载Excel知识图谱数据"""
        print(f"加载知识图谱Excel文件: {excel_path}")
        
        try:
            # 读取Excel文件
            data = pd.read_excel(excel_path, sheet_name='Sheet1', header=None)
            df = data.iloc[1:].copy()
            
            # 设置列名
            df.columns = df.iloc[0].tolist()
            df = df.iloc[1:].copy()   
            # 确保需要填充的列是字符串类型
            fill_columns = ['一级', '二级', '三级']
            for col in fill_columns:
                df[col] = df[col].astype(str)  # 先转为字符串
                df[col] = df[col].replace('nan', np.nan)  # 将字符串'nan'转为真正的NaN
                df[col] = df[col].replace('None', np.nan)  # 将字符串'None'转为真正的NaN
                df[col] = df[col].fillna(method='ffill')  # 前向填充
            
            print("数据填充完成!")

            knowledge_docs = []
            for index, row in df.iterrows():
                if pd.isna(row['一级']) or row['一级'] == '一级':
                    continue

                doc_content = self._build_knowledge_document(row)
                if doc_content:
                    # 创建Document对象
                    doc = Document(
                        text=doc_content,
                        metadata={
                            'file_name': '扬州市人工智能产业图谱.xlsx',
                            'folder_type': 'knowledge_graph',
                            'doc_id': f"kg_{index}",
                            'title': f"{row.get('企业', '')} - {row.get('一级', '')}",
                            'industry_level_1': row.get('一级', ''),
                            'industry_level_2': row.get('二级', ''),
                            'industry_level_3': row.get('三级', ''),
                            'company': row.get('企业', ''),
                            'business': row.get('涉及业务', ''),
                            'region': row.get('地区', ''),
                            'source_type': 'knowledge_graph'
                        }
                    )
                    knowledge_docs.append(doc)
            
            print(f"✓ 知识图谱数据加载完成: {len(knowledge_docs)} 个企业节点")
            return knowledge_docs
            
        except Exception as e:
            print(f"✗ 加载知识图谱Excel失败: {e}")
            return []
    
    def _build_knowledge_document(self, row) -> str:
        """构建知识图谱文档内容"""
        parts = []
        
        # 基本信息
        if pd.notna(row.get('企业')) and row['企业']:
            parts.append(f"企业名称：{row['企业']}")
        
        if pd.notna(row.get('一级')) and row['一级']:
            parts.append(f"产业层级：{row['一级']}")
        
        if pd.notna(row.get('二级')) and row['二级']:
            parts.append(f"细分领域：{row['二级']}")
        
        if pd.notna(row.get('三级')) and row['三级']:
            parts.append(f"具体分类：{row['三级']}")
        
        if pd.notna(row.get('涉及业务')) and row['涉及业务']:
            parts.append(f"主营业务：{row['涉及业务']}")
        
        if pd.notna(row.get('地区')) and row['地区']:
            parts.append(f"所在地区：{row['地区']}")
        
        # 构建完整描述
        if parts:
            return "\n".join(parts)
        return None
    
    def load_research_institutions(self, excel_path: str) -> List[Document]:
        """加载研究机构数据（Sheet2）"""
        try:
            data = pd.read_excel(excel_path, sheet_name='Sheet2', header=None)
            df = data.copy()
            
            df.columns = ['region', 'institution']
            research_docs = []
            
            for index, row in df.iterrows():
                
                region = row['region']
                institution = row['institution']
                
                if institution:
                    doc_content = f"地区：{region}\n研究机构：{institution}\n类型：人工智能研究机构"
                    
                    doc = Document(
                        text=doc_content,
                        metadata={
                            'file_name': '扬州市人工智能产业图谱.xlsx',
                            'folder_type': 'knowledge_graph',
                            'doc_id': f"research_{index}",
                            'title': institution,
                            'region': region,
                            'institution': institution,
                            'source_type': 'research_institution'
                        }
                    )
                    research_docs.append(doc)
            print(f"✓ 研究机构数据加载完成: {len(research_docs)} 个机构")
            return research_docs
            
        except Exception as e:
            print(f"✗ 加载研究机构数据失败: {e}")
            return []

class QueryRewriter:
    """查询改写器 - 生成同义问题扩展检索"""
    
    def __init__(self, llm):
        self.llm = llm
    
    def rewrite_queries(self, original_query: str, num_queries: int = 2) -> List[str]:
        """生成同义查询"""
        prompt = f"""
请为以下问题生成{num_queries}个不同但语义相似的查询问题。这些查询应该从不同角度表达相同的意思，以帮助检索系统找到更全面的相关信息。

原始问题：{original_query}

要求：
1. 保持核心语义不变
2. 使用不同的表达方式和角度
3. 涵盖政策文件可能使用的不同术语
4. 每个查询都应该是完整的问题

请直接输出{num_queries}个查询，每个查询一行：
"""
        messages = [ChatMessage(role="user", content=prompt)]
        try:
            response = self.llm.chat(messages)
            pattern = r'(.*?)</think>(.*)'
            match = re.search(pattern, response.content, re.DOTALL)
            if match:
                content = match.group(2).strip()
            else:
                content = response.content.strip()
        
            queries = [q.strip() for q in content.split('\n') if q.strip()]
            # 确保包含原始查询
            if original_query not in queries:
                queries = [original_query] + queries
            return queries[:num_queries]
        except Exception as e:
            print(f"查询改写失败: {e}")
            return [original_query]

class QueryDecomposer:
    """查询分解器 - 将复杂查询拆分为子问题"""
    
    def __init__(self, llm):
        self.llm = llm
    
    def decompose_query(self, complex_query: str) -> List[str]:
        """分解复杂查询为子问题"""
        prompt = f"""
请将以下复杂查询分解为2-4个更简单、更具体的子问题。这些子问题应该涵盖原查询的各个方面，便于分别检索相关信息。

复杂查询：{complex_query}

要求：
1. 每个子问题应该独立且具体
2. 子问题之间应该有逻辑关联
3. 涵盖原查询的所有关键方面
4. 保持政策查询的专业性

请直接输出子问题，每个问题一行：
"""
        messages = [ChatMessage(role="user", content=prompt)]
        try:
            response = self.llm.chat(messages)
            pattern = r'(.*?)</think>(.*)'
            match = re.search(pattern, response.content, re.DOTALL)
            if match:
                content = match.group(2).strip()
            else:
                content = response.content.strip()
                
            sub_queries = [q.strip() for q in content.split('\n') if q.strip()]
            return sub_queries if sub_queries else [complex_query]
        except Exception as e:
            print(f"查询分解失败: {e}")
            return [complex_query]

class HypotheticalAnswerGenerator:
    """假设答案生成器 - 生成理想答案用于检索"""
    
    def __init__(self, llm):
        self.llm = llm
    
    def generate_hypothetical_answer(self, query: str) -> str:
        """为查询生成假设答案"""
        prompt = f"""
针对以下政策查询问题，请生成一个理想的、全面的答案。这个答案将用于检索相关的政策文档。

查询：{query}

要求：
1. 想象一个完美的政策答案应该包含哪些内容
2. 涵盖查询的所有方面
3. 使用政策文档中可能出现的专业术语
4. 保持客观、准确的政策语言风格

请生成假设答案：
"""
        messages = [ChatMessage(role="user", content=prompt)]
        
        try:
            response = self.llm.chat(messages)

            pattern = r'(.*?)</think>(.*)'
            match = re.search(pattern, response.content, re.DOTALL)
            if match:
                content = match.group(2).strip()
            else:
                content = response.content.strip()

            return content.strip()
        except Exception as e:
            print(f"假设答案生成失败: {e}")
            return query

class ParagraphSplitter:
    """段落分割器 - 基于语义段落而非固定长度"""
    
    def __init__(self, max_length: int = 1000, min_length: int = 50):
        self.max_length = max_length
        self.min_length = min_length
    
    def split_document(self, document: Document) -> List[Document]:
        """将文档分割为段落"""
        # 对于知识图谱数据，不进行分割
        if document.metadata.get('source_type') in ['knowledge_graph', 'research_institution']:
            return [document]
            
        text = document.text
        paragraphs = []
        
        # 按段落分隔符分割
        raw_paragraphs = text.split(EnhancedConfig.PARAGRAPH_SEPARATOR)
        
        for para in raw_paragraphs:
            para = para.strip()
            if not para:
                continue
                
            # 如果段落过长，按句子进一步分割
            if len(para) > self.max_length:
                sentences = re.split(r'[。！？!?]', para)
                current_chunk = ""
                
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                        
                    if len(current_chunk) + len(sentence) <= self.max_length:
                        current_chunk += sentence + "。"
                    else:
                        if current_chunk and len(current_chunk) >= self.min_length:
                            paragraphs.append(current_chunk.strip())
                        current_chunk = sentence + "。"
                
                if current_chunk and len(current_chunk) >= self.min_length:
                    paragraphs.append(current_chunk.strip())
            else:
                if len(para) >= self.min_length:
                    paragraphs.append(para)
        
        # 创建新的Document对象
        paragraph_docs = []
        for i, para_text in enumerate(paragraphs):
            new_doc = Document(
                text=para_text,
                metadata=document.metadata.copy()
            )
            new_doc.metadata['paragraph_id'] = i
            new_doc.metadata['chunk_type'] = 'paragraph'
            paragraph_docs.append(new_doc)
        
        return paragraph_docs

class SentenceChunkSplitter:
    """句子分块器 - 基于LlamaIndex的SentenceSplitter"""
    
    def __init__(self, chunk_size: int = 1024, chunk_overlap: int = 200, min_chunk_size: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator="\n"  # 按换行符拆分，保留段落结构
        )
    
    def split_document(self, document: Document) -> List[Document]:
        """将文档分割为句子块"""
        # 对于知识图谱数据，不进行分割
        if document.metadata.get('source_type') in ['knowledge_graph', 'research_institution']:
            return [document]
            
        # 使用LlamaIndex的SentenceSplitter，这里split_text返回的是字符串列表
        text_chunks = self.splitter.split_text(document.text)
        
        # 转换为Document对象
        chunk_docs = []
        for i, chunk_text in enumerate(text_chunks):
            if len(chunk_text) < self.min_chunk_size:
                continue
                
            new_doc = Document(
                text=chunk_text,
                metadata=document.metadata.copy()
            )
            new_doc.metadata['chunk_id'] = i
            new_doc.metadata['chunk_type'] = 'sentence'
            chunk_docs.append(new_doc)
        
        return chunk_docs

class DocumentSplitterFactory:
    """文档分割器工厂 - 根据配置选择分割方法"""
    
    @staticmethod
    def create_splitter() -> object:
        """创建分割器实例"""
        if EnhancedConfig.CHUNK_MODE == 'paragraph':
            return ParagraphSplitter(
                max_length=EnhancedConfig.MAX_PARAGRAPH_LENGTH,
                min_length=EnhancedConfig.MIN_PARAGRAPH_LENGTH
            )
        elif EnhancedConfig.CHUNK_MODE == 'sentence':
            return SentenceChunkSplitter(
                chunk_size=EnhancedConfig.CHUNK_SIZE,
                chunk_overlap=EnhancedConfig.CHUNK_OVERLAP,
                min_chunk_size=EnhancedConfig.MIN_CHUNK_SIZE
            )
        else:
            raise ValueError(f"不支持的CHUNK_MODE: {EnhancedConfig.CHUNK_MODE}")

class ReciprocalRankFusion:
    """倒数排序融合算法"""
    
    def __init__(self, k: int = 60):
        self.k = k
    
    def fuse(self, ranked_lists: List[List[NodeWithScore]]) -> List[NodeWithScore]:
        """融合多个排序列表"""
        if not ranked_lists:
            return []
        
        # 初始化分数字典
        scores = {}
        
        # 对每个排序列表计算RRF分数
        for rank_list in ranked_lists:
            for rank, node in enumerate(rank_list):
                node_id = node.node.node_id
                if node_id not in scores:
                    scores[node_id] = 0.0
                scores[node_id] += 1.0 / (self.k + rank + 1)
        
        # 创建融合后的节点列表
        fused_nodes = []
        for node_id, score in scores.items():
            # 找到对应的节点（取第一个出现的位置）
            for rank_list in ranked_lists:
                for node in rank_list:
                    if node.node.node_id == node_id:
                        fused_nodes.append(NodeWithScore(
                            node=node.node,
                            score=score
                        ))
                        break
                else:
                    continue
                break

        # 按分数降序排序
        fused_nodes.sort(key=lambda x: x.score, reverse=True)
        return fused_nodes

class SimHashDeduplicator:
    """SimHash去重器"""

    def __init__(self, config: EnhancedConfig):
        self.config = config
        self.hash_cache: Dict[str, int] = {}
        self.hash_set: Set[int] = set()
        
    def _tokenize(self, text: str) -> List[str]:
        """文本分词"""
        text = re.sub(r'[^\w\u4e00-\u9fa5]', ' ', text)
        words = jieba.cut(text)
        stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这', '那', '他', '她', '它'}
        tokens = [word for word in words if len(word) > 1 and word not in stop_words]
        return tokens
    
    def _get_word_hash(self, word: str) -> int:
        """获取词语哈希值"""
        md5 = hashlib.md5(word.encode('utf-8'))
        return int(md5.hexdigest()[:16], 16)
    
    def _simhash(self, text: str) -> int:
        """计算SimHash值"""
        if text in self.hash_cache:
            return self.hash_cache[text]
            
        tokens = self._tokenize(text)
        if not tokens:
            return 0
            
        vector = [0] * self.config.HASH_BITS
        
        for token in tokens:
            token_hash = self._get_word_hash(token)
            for i in range(self.config.HASH_BITS):
                bit_mask = 1 << i
                if token_hash & bit_mask:
                    vector[i] += 1
                else:
                    vector[i] -= 1
        
        fingerprint = 0
        for i in range(self.config.HASH_BITS):
            if vector[i] > 0:
                fingerprint |= 1 << i
                
        self.hash_cache[text] = fingerprint
        return fingerprint

    def _hamming_distance(self, hash1: int, hash2: int) -> int:
        """计算汉明距离"""
        xor_result = hash1 ^ hash2
        distance = 0
        while xor_result:
            distance += 1
            xor_result &= xor_result - 1
        return distance
    
    def deduplicate_nodes(self, nodes: List[NodeWithScore]) -> List[NodeWithScore]:
        """节点去重"""
        unique_nodes = []
        seen_hashes = set()
        
        for node in nodes:
            node_text = node.node.text
            if len(node_text) < self.config.MIN_TEXT_LENGTH:
                unique_nodes.append(node)
                continue
                
            node_hash = self._simhash(node_text)
            is_duplicate = False
            
            for existing_hash in seen_hashes:
                if self._hamming_distance(node_hash, existing_hash) <= self.config.SIMILAR_THRESHOLD:
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                unique_nodes.append(node)
                seen_hashes.add(node_hash)
        
        return unique_nodes

class KeyTermExtractor:
    """关键信息提取器 - 优化版本"""
    
    @staticmethod
    def extract_key_terms(query: str) -> Dict[str, Any]:
        """提取关键信息 - 增强版本"""
        key_terms = {
            'years': [],
            'numbers': [],
            'policy_names': [],
            'locations': [],
            'key_entities': [],
            'time_range': None,
            'key_phrases': []  # 新增：关键短语
        }

        # 提取年份（4位数字）
        years = re.findall(r'(?<!\d)(19\d{2}|20\d{2})(?!\d)', query)
        key_terms['years'] = list(set(years))

        # 提取时间范围（增强模式）
        time_range_patterns = [
            r'(\b(19|20)\d{2})\s*[-至]\s*(\b(19|20)\d{2})',  # 2023-2025 或 2023至2025
            r'(\b(19|20)\d{2})年\s*(?:[-至]\s*)?(\b(19|20)\d{2})年',  # 2023年-2025年
        ]
        
        for pattern in time_range_patterns:
            time_range_match = re.search(pattern, query)
            if time_range_match:
                start_year = int(time_range_match.group(1))
                end_year = int(time_range_match.group(3))
                key_terms['time_range'] = {'start': start_year, 'end': end_year}
                break
        
        # 提取数字（包括小数和百分比）
        numbers = re.findall(r'\b\d+(?:\.\d+)?%?\b', query)
        key_terms['numbers'] = list(set(numbers))
        
        # 提取政策名称（书名号内内容）
        policy_names = re.findall(r'《([^《》]+)》', query)
        key_terms['policy_names'] = policy_names
        
        # 提取地点（扬州相关 - 增强）
        locations = re.findall(r'(江都区|高邮市|仪征市|广陵区|邗江区|宝应县|景区|经济开发区|生态科技新城|经济技术开发区|开发区|扬州)', query)
        key_terms['locations'] = list(set(locations))
        
        # 提取关键实体
        entities = re.findall(r'(任务清单|主要目标|实施方案|政策|通知|办法|条例|规定|细则|指南|规划|计划|方案|工作要点|重点工作)', query)
        key_terms['key_entities'] = list(set(entities))
        
        # 提取关键短语
        phrases = re.findall(r'(?:^|\s)([^，。！？\s]{2,6}?(?:任务|目标|计划|方案|政策|措施|工作))', query)
        key_terms['key_phrases'] = phrases
        
        return key_terms
    
    @staticmethod
    def calculate_semantic_similarity(file_name: str, query: str) -> float:
        """计算文档名称与查询的语义相似度"""
        # 简单的文本相似度计算，可以替换为更复杂的模型
        file_words = set(re.findall(r'[\u4e00-\u9fff]+', file_name))
        query_words = set(re.findall(r'[\u4e00-\u9fff]+', query))
        
        if not file_words or not query_words:
            return 0.0
            
        intersection = file_words & query_words
        union = file_words | query_words
        
        return len(intersection) / len(union) if union else 0.0
    
    @staticmethod
    def calculate_key_term_boost(file_name: str, key_terms: Dict[str, Any], query: str) -> Tuple[float, Dict[str, float]]:
        """计算文档名称的关键信息匹配度提升 - 返回详细的分项得分"""
        boost = 0.0
        penalty = 0.0
        detailed_scores = {
            'semantic_similarity': 0.0,
            'year_match': 0.0,
            'policy_match': 0.0,
            'location_match': 0.0,
            'phrase_match': 0.0,
            'time_range_match': 0.0,
            'location_penalty': 0.0,
            'time_penalty': 0.0,
            'year_mismatch_penalty': 0.0
        }
        
        # 1. 语义相似度基础分 (权重最高)
        semantic_similarity = KeyTermExtractor.calculate_semantic_similarity(file_name, query)
        semantic_boost = semantic_similarity * 0.4  # 提高语义相似度权重
        boost += semantic_boost
        detailed_scores['semantic_similarity'] = semantic_boost

        # 提取文档名称中的年份
        file_years = re.findall(r'(?<!\d)(19\d{2}|20\d{2})(?!\d)', file_name)
        file_years_int = [int(year) for year in file_years] if file_years else []

        # 2. 年份匹配 (分级权重)
        year_boost = 0.0
        year_mismatch_penalty = 0.0

        if key_terms['years']:  # 查询中存在年份
            matched_years = []
            for year in key_terms['years']:
                if year in file_name:
                    matched_years.append(year)
                    if file_name.startswith(year) or f"{year}年" in file_name:
                        year_boost += 1.0  # 重要位置年份
                    else:
                        year_boost += 1.0  # 普通位置年份

            # 年份不匹配惩罚：查询中有年份但文档名称中的年份都不匹配
            if not matched_years and file_years:  # 文档有年份但与查询不匹配
                year_mismatch_penalty += 3.0
                print(f"  ✗ 年份不匹配: 查询指定年份{key_terms['years']}，但文档名称年份{file_years}不匹配, -2")

            # elif not matched_years and key_terms['years']:  # 查询有年份但文档无年份
            #     year_mismatch_penalty += 0.1
            #     print(f"  ✗ 年份缺失: 查询指定年份{key_terms['years']}，但文档名称无年份信息, -0.4")

        if year_boost > 0:
            year_boost = min(year_boost, 2.0)  # 限制最大奖励
            boost += year_boost
            detailed_scores['year_match'] = year_boost

        # 3. 政策名称匹配
        policy_boost = 0.0
        for policy_name in key_terms['policy_names']:
            if policy_name in file_name:
                policy_boost += 0.5

        if policy_boost > 0:
            policy_boost = min(policy_boost, 1.8)
            boost += policy_boost
            detailed_scores['policy_match'] = policy_boost

        # 4. 地点匹配
        location_boost = 0.0
        for location in key_terms['locations']:
            if location in file_name:
                location_boost += 0.3

        if location_boost > 0:
            location_boost = min(location_boost, 0.9)
            boost += location_boost
            detailed_scores['location_match'] = location_boost

        # 5. 关键短语匹配
        phrase_boost = 0.0
        for phrase in key_terms['key_phrases']:
            if phrase in file_name:
                phrase_boost += 0.4

        if phrase_boost > 0:
            phrase_boost = min(phrase_boost, 1.2)
            boost += phrase_boost
            detailed_scores['phrase_match'] = phrase_boost

        # 6. 时间范围匹配 (优化逻辑)
        time_range_boost = 0.0
        time_penalty = 0.0

        if key_terms['time_range']:
            start_year = key_terms['time_range']['start']
            end_year = key_terms['time_range']['end']

            if file_years_int:
                max_file_year = max(file_years_int) if file_years_int else 0

                # 精确匹配奖励
                in_range_years = [year for year in file_years_int if start_year <= year <= end_year]
                if in_range_years:
                    time_range_boost += 0.8
                else:
                    # 相邻年份奖励 (±1-2年)
                    adjacent_years = [
                        year for year in file_years_int 
                        if start_year - 1 <= year <= end_year + 1
                    ]
                    if adjacent_years:
                        time_range_boost += 0.1
                    else:
                        # 时间完全不相关惩罚
                        if max_file_year < start_year - 5:  # 过于陈旧
                            time_penalty += 3.0
                        elif max_file_year > end_year + 5:  # 过于超前
                            time_penalty += 3.0

        boost += time_range_boost
        penalty += time_penalty
        detailed_scores['time_range_match'] = time_range_boost
        detailed_scores['time_penalty'] = time_penalty

        # 7. 地点不匹配惩罚
        location_penalty = 0.0
        if key_terms['locations']:
            location_matched = any(location in file_name for location in key_terms['locations'])
            if not location_matched:
                # 检查是否是同省份其他城市
                jiangsu_cities = ['南京', '南京市', '苏州', '苏州市', '无锡', '无锡市', '常州', '常州市', 
                                 '镇江', '镇江市', '南通', '南通市', '泰州', '泰州市', '徐州', '徐州市', 
                                 '宿迁', '宿迁市', '连云港', '连云港市', '淮安', '淮安市', '盐城', '盐城市']
                file_other_locations = [city for city in jiangsu_cities if city in file_name]

                if file_other_locations:
                    location_penalty += 3.0 # 同省份其他城市，轻度惩罚
                else:
                    location_penalty += 3.0  # 完全不同的地点，中度惩罚

        penalty += location_penalty
        detailed_scores['location_penalty'] = location_penalty

        # 8. 年份不匹配惩罚
        penalty += year_mismatch_penalty
        detailed_scores['year_mismatch_penalty'] = year_mismatch_penalty

        # 计算最终调整分数
        final_adjustment = boost - penalty

        # 限制调整范围在合理区间 [-1.0, 2.0]
        final_adjustment = max(-3.0, min(3.0, final_adjustment))

        return final_adjustment, detailed_scores
    
    @staticmethod
    def apply_key_term_adjustment(nodes: List[NodeWithScore], query: str) -> List[NodeWithScore]:
        """应用关键信息调整到节点列表 - 优化版本"""
        if not nodes:
            return nodes
            
        key_terms = KeyTermExtractor.extract_key_terms(query)
        
        # 打印提取的关键信息
        print(f"\n=== 关键信息提取 ===")
        print(f"查询: {query}")
        for category, terms in key_terms.items():
            if terms or (category == 'time_range' and terms is not None):
                print(f"{category}: {terms}")
        
        print(f"\n=== 关键信息调整 ===")
        
        adjusted_nodes = []
        adjustment_details = []
        
        for i, node in enumerate(nodes):
            original_score = node.score
            file_name = node.node.metadata.get('file_name', '')
            adjustment, detailed_scores = KeyTermExtractor.calculate_key_term_boost(file_name, key_terms, query)
            
            # 智能评分融合策略
            new_score = KeyTermExtractor._intelligent_score_fusion(original_score, adjustment, detailed_scores, file_name)
            
            # 创建新节点
            adjusted_node = NodeWithScore(
                node=node.node,
                score=new_score
            )
            adjusted_nodes.append(adjusted_node)
            adjustment_details.append((i, original_score, new_score, adjustment))
            
            if adjustment != 0:
                print(f"节点 {i+1}: 原始分数 {original_score:.3f} -> 调整后 {new_score:.3f} (调整: {adjustment:+.2f})")
        
        # 重新排序
        adjusted_nodes.sort(key=lambda x: x.score, reverse=True)
        
        # 打印调整统计
        if nodes:
            original_top_score = nodes[0].score
            adjusted_top_score = adjusted_nodes[0].score
            score_change = adjusted_top_score - original_top_score

            print(f"\n=== 调整统计 ===")
            print(f"最高分: {original_top_score:.3f} -> {adjusted_top_score:.3f} ({score_change:+.3f})")

            # 显示分数分布
            original_scores = [node.score for node in nodes[:5]]
            adjusted_scores = [node.score for node in adjusted_nodes[:5]]
            print(f"前5名原始分数: {[f'{s:.3f}' for s in original_scores]}")
            print(f"前5名调整分数: {[f'{s:.3f}' for s in adjusted_scores]}")

        return adjusted_nodes

    @staticmethod
    def _intelligent_score_fusion(original_score: float, total_boost: float, 
                                detailed_scores: Dict[str, float], file_name: str) -> float:
        """智能评分融合策略"""
        
        # 基础调整：使用sigmoid函数进行平滑调整
        def sigmoid_adjustment(x):
            return 4 / (1 + math.exp(-2 * x)) - 2  # 输出范围[-2, 2]
        
        # 1. 对于高原始分数，调整要更谨慎
        if original_score > 1.5:
            adjustment_factor = 0.3  # 高分数文档调整幅度较小
        elif original_score > 0.8:
            adjustment_factor = 0.6  # 中等分数文档调整幅度中等
        else:
            adjustment_factor = 0.8  # 低分数文档调整幅度较大
        
        # 2. 使用sigmoid函数平滑调整值
        smoothed_boost = sigmoid_adjustment(total_boost) * adjustment_factor
        
        # 3. 考虑各项得分的权重分布
        positive_components = sum([
            detailed_scores['semantic_similarity'],
            detailed_scores['year_match'],
            detailed_scores['policy_match'],
            detailed_scores['location_match'],
            detailed_scores['phrase_match'],
            detailed_scores['time_range_match']
        ])
        
        negative_components = sum([
            detailed_scores['location_penalty'],
            detailed_scores['time_penalty']
        ])
        
        # 4. 计算置信度权重
        confidence_weight = min(1.0, positive_components / 3.0)  # 正项得分越高，置信度越高
        
        # 5. 最终调整计算
        if total_boost > 0:
            # 正向调整：基于置信度加权
            final_adjustment = smoothed_boost * confidence_weight
            new_score = original_score * (1 + final_adjustment)
        else:
            # 负向调整：更谨慎，考虑原始分数
            penalty_severity = 0.5 if original_score > 0.6 else 0.3
            final_adjustment = smoothed_boost * penalty_severity
            new_score = original_score * (1 + final_adjustment)
        
        # 确保分数在合理范围内 [0, 1]
        new_score = max(0.0, min(1.0, new_score))
        
        return new_score

    @staticmethod
    def _print_adjustment_details(file_name: str, original_score: float, 
                                new_score: float, total_boost: float, 
                                detailed_scores: Dict[str, float]):
        """打印调整详情"""
        print(f"\n📄 文档: {file_name}")
        print(f"   原始分数: {original_score:.3f} → 新分数: {new_score:.3f}")
        print(f"   总调整值: {total_boost:.2f}")
        
        # 打印正项得分
        positive_scores = {k: v for k, v in detailed_scores.items() if v > 0 and 'penalty' not in k}
        if positive_scores:
            print("   ✅ 奖励项:")
            for key, value in positive_scores.items():
                if value > 0:
                    print(f"     - {key}: +{value:.2f}")
        
        # 打印负项得分
        negative_scores = {k: v for k, v in detailed_scores.items() if v > 0 and 'penalty' in k}
        if negative_scores:
            print("   ❌ 惩罚项:")
            for key, value in negative_scores.items():
                if value > 0:
                    print(f"     - {key}: -{value:.2f}")

    @staticmethod
    def _print_fusion_statistics(adjustment_stats: List[dict], filtered_nodes: List[NodeWithScore]):
        """打印融合统计信息"""
        print(f"\n=== 智能评分融合统计 ===")
        print(f"处理节点总数: {len(adjustment_stats)}")
        
        if adjustment_stats:
            avg_original = sum(stat['original_score'] for stat in adjustment_stats) / len(adjustment_stats)
            avg_new = sum(stat['new_score'] for stat in adjustment_stats) / len(adjustment_stats)
            avg_adjustment = sum(stat['adjustment'] for stat in adjustment_stats) / len(adjustment_stats)
            
            print(f"平均原始分数: {avg_original:.3f}")
            print(f"平均新分数: {avg_new:.3f}")
            print(f"平均调整幅度: {avg_adjustment:+.3f}")
            
            # 调整幅度分布
            positive_adjustments = [stat for stat in adjustment_stats if stat['adjustment'] > 0]
            negative_adjustments = [stat for stat in adjustment_stats if stat['adjustment'] < 0]
            no_adjustments = [stat for stat in adjustment_stats if stat['adjustment'] == 0]
            
            print(f"分数提升文档: {len(positive_adjustments)}个")
            print(f"分数降低文档: {len(negative_adjustments)}个")
            print(f"分数不变文档: {len(no_adjustments)}个")
        
        if filtered_nodes:
            print(f"调整后最高分: {filtered_nodes[0].score:.3f}")
            print(f"调整后最低分: {filtered_nodes[-1].score:.3f}")

class Qwen3Embedding:
    """Qwen3 Embedding模型封装"""

    def __init__(self, model_name: str = "Qwen/Qwen3-Embedding-8B", instruction: str = ""):
        self.model_name = model_name
        self.instruction = instruction
        self.model = None
        self.tokenizer = None
        self._initialize_model()

    def _initialize_model(self):
        """初始化模型"""
        print(f"加载Qwen3 Embedding模型: {self.model_name}")
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name, trust_remote_code=True)
        self.model = AutoModel.from_pretrained(
            self.model_name,
            trust_remote_code=True,
            torch_dtype=torch.float16,
            device_map=EnhancedConfig.DEVICE
        )
        self.model.eval()
        print("Qwen3 Embedding模型加载成功")

    def get_query_embedding(self, query: str) -> List[float]:
        """获取查询嵌入"""
        if self.instruction:
            query = f"{self.instruction}\n{query}"

        with torch.no_grad():
            inputs = self.tokenizer(
                query,
                padding=True,
                truncation=True,
                return_tensors="pt",
                max_length=8192
            ).to(EnhancedConfig.DEVICE)
    
            outputs = self.model(**inputs)
            # 使用平均池化获取嵌入
            embeddings = self._mean_pooling(outputs.last_hidden_state, inputs['attention_mask'])
            return embeddings[0].cpu().numpy().tolist()

    def get_text_embeddings(self, texts: List[str]) -> List[List[float]]:
        """获取文本嵌入列表"""
        if self.instruction:
            texts = [f"{self.instruction}\n{text}" for text in texts]

        with torch.no_grad():
            inputs = self.tokenizer(
                texts,
                padding=True,
                truncation=True,
                return_tensors="pt",
                max_length=8192
            ).to(EnhancedConfig.DEVICE)

            outputs = self.model(**inputs)
            # 使用平均池化获取嵌入
            embeddings = self._mean_pooling(outputs.last_hidden_state, inputs['attention_mask'])
            return embeddings.cpu().numpy().tolist()

    def _mean_pooling(self, model_output, attention_mask):
        """平均池化获取句子嵌入"""
        token_embeddings = model_output
        input_mask_expanded = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
        return torch.sum(token_embeddings * input_mask_expanded, 1) / torch.clamp(
            input_mask_expanded.sum(1), min=1e-9)

class Qwen3Reranker:
    """Qwen3 Reranker模型封装"""

    def __init__(self, model_name: str = "Qwen/Qwen3-Reranker-8B", instruction: str = ""):
        self.model_name = model_name
        self.instruction = instruction
        self.model = None
        self.tokenizer = None
        self._initialize_model()

    def _initialize_model(self):
        """初始化模型"""

        print(f"加载Qwen3 Reranker模型: {self.model_name}")

        self.tokenizer = AutoTokenizer.from_pretrained(
            self.model_name,
            trust_remote_code=True)
        
        self.model = AutoModelForCausalLM.from_pretrained(
            self.model_name,
            trust_remote_code=True,
            torch_dtype=torch.float16,
            device_map=EnhancedConfig.DEVICE)
        
        self.model.eval()
        print("Qwen3 Reranker模型加载成功")

    def compute_score(self, pairs: List[Tuple[str, str]]) -> List[float]:
        """计算查询-文档对的相关性分数"""
        scores = []
        for query, document in pairs:

            q_d_data = f"查询: {query}\n文档: {document}"
            messages = []
            messages.append({"role": "system", "content":self.instruction})
            messages.append({"role": "assistant","content":q_d_data})
            
            messages_inputs = self.tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True, return_tensors="pt")
            
            inputs = self.tokenizer(messages_inputs, return_tensors="pt").to(EnhancedConfig.DEVICE)
            with torch.no_grad():
                outputs = self.model(**inputs)
                logits = outputs.logits[:, -1, :]  # 取最后一个token的logits

                yes_token_id = self.tokenizer.encode("yes", add_special_tokens=False)[0]
                no_token_id = self.tokenizer.encode("no", add_special_tokens=False)[0]

                yes_logit = logits[0, yes_token_id].item()
                no_logit = logits[0, no_token_id].item()
                exp_yes = np.exp(yes_logit)
                exp_no = np.exp(no_logit)
                score = exp_yes / (exp_yes + exp_no)
                scores.append(score)

        return scores

class SimpleLLM:
    """简化的LLM包装器，不使用LlamaIndex的抽象基类"""
    def __init__(self, model_path):

        self.tokenizer = AutoTokenizer.from_pretrained(model_path)
        self.llm = LLM(model=model_path, gpu_memory_utilization=EnhancedConfig.gpu_memory_utilization, max_model_len=EnhancedConfig.max_model_len)
        self.sampling_params = SamplingParams(
            temperature=EnhancedConfig.TEMPERATURE,
            max_tokens=EnhancedConfig.MAX_NEW_TOKENS,
            top_p=EnhancedConfig.TOP_P)
    
    def chat(self, messages: List[ChatMessage]) -> ChatMessage:
        """聊天接口"""
        # 转换消息格式：LlamaIndex的ChatMessage -> OpenAI API格式
        openai_messages = []
        for msg in messages:
            openai_messages.append({
                "role": msg.role,
                "content": msg.content})

        prompts = self.tokenizer.apply_chat_template(openai_messages, 
                                                tokenize=False, 
                                                add_generation_prompt=True, 
                                                return_tensors="pt")
        response = self.llm.generate(prompts,self.sampling_params)

        return ChatMessage(
            role="assistant",
            content=response[0].outputs[0].text)

class EnhancedDocumentProcessor:
    """增强文档处理器 - 支持多格式文档读取和知识图谱数据"""
    
    def __init__(self):
        self.documents = []
        self.knowledge_processor = KnowledgeGraphProcessor()

    def read_all_documents(self, base_path: str, knowledge_graph_path: str = None, include_knowledge_graph: bool = True) -> List[Document]:
        """读取所有文档，可选择是否包含知识图谱数据"""
        all_docs = []
        
        # 1. 读取传统文档
        print("=== 读取传统文档 ===")
        traditional_docs = self._read_traditional_documents(base_path)
        all_docs.extend(traditional_docs)
        
        # 2. 读取知识图谱数据（可选）
        if include_knowledge_graph and knowledge_graph_path and os.path.exists(knowledge_graph_path):
            print("\n=== 读取知识图谱数据 ===")
            knowledge_docs = self.knowledge_processor.load_excel_knowledge_graph(knowledge_graph_path)
            all_docs.extend(knowledge_docs)
            
            research_docs = self.knowledge_processor.load_research_institutions(knowledge_graph_path)
            all_docs.extend(research_docs)
            print(f"✓ 知识图谱数据已加载")
        elif include_knowledge_graph and knowledge_graph_path:
            print(f"\n⚠ 知识图谱文件不存在: {knowledge_graph_path}")
        else:
            print(f"\n⚠ 知识图谱数据加载已禁用")
        
        print(f"\n总计读取 {len(all_docs)} 个文档")
        print(f"  - 传统文档: {len(traditional_docs)}")
        if include_knowledge_graph:
            print(f"  - 知识图谱企业: {len([d for d in all_docs if d.metadata.get('source_type') == 'knowledge_graph'])}")
            print(f"  - 研究机构: {len([d for d in all_docs if d.metadata.get('source_type') == 'research_institution'])}")
        
        return all_docs

    def _read_traditional_documents(self, base_path: str) -> List[Document]:
        """读取传统文档（原有逻辑）"""
        folders = {
            "人工智能产业链招商": "ai_industry",
            "扬州公积金政策": "housing_fund", 
            "扬州人社局相关政策": "hr_policy",
            "扬州政务政策解读": "government_policy"
        }

        all_docs = []
        doc_id = 0

        for folder_name, folder_type in folders.items():
            folder_path = os.path.join(base_path, folder_name)
            
            if not os.path.exists(folder_path):
                print(f"警告: 文件夹不存在 {folder_path}")
                continue

            for file in os.listdir(folder_path):
                file_path = os.path.join(folder_path, file)
                try:
                    if file.endswith('.docx'):
                        content = self._read_docx(file_path)
                    elif file.endswith('.txt'):
                        content = self._read_txt(file_path)
                    else:
                        continue

                    if not content.strip():
                        continue

                    content = self._clean_text(content)
                    title = self._extract_title(content)

                    doc = Document(
                        text=content,
                        metadata={
                            'file_name': file,
                            'folder_type': folder_type,
                            'doc_id': doc_id,
                            'title': title,
                            'source_type': 'traditional_doc'
                        }
                    )
                    all_docs.append(doc)
                    doc_id += 1
                    print(f"  ✓ {file} ({len(content)} 字符)")

                except Exception as e:
                    print(f"  ✗ {file}: {str(e)}")

        return all_docs

    # 保留原有的辅助方法
    def _read_docx(self, file_path: str) -> str:
        """读取docx格式文件内容"""
        doc = docx.Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return '\n'.join(paragraphs)

    def _read_txt(self, file_path: str) -> str:
        """读取txt格式文件内容"""
        encodings = ['utf-8', 'gbk', 'gb2312']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except Exception:
                continue
        return ""

    def _clean_text(self, text: str) -> str:
        """文本清洗"""
        text = re.sub(r'\r\n|\r', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        return text.strip()

    def _extract_title(self, text: str) -> str:
        """从文本中提取标题"""
        lines = text.split('\n')
        policy_keywords = ['通知', '意见', '办法', '方案', '条例', '规定', '指南', '细则']

        for line in lines[:5]:
            line = line.strip()
            if 5 < len(line) < 100:
                if any(keyword in line for keyword in policy_keywords):
                    return line
        
        return lines[0][:50].strip() if lines else "未命名文档"

class Qwen3DirectReranker:
    """Qwen3 Reranker直接封装 - 适配LlamaIndex NodeWithScore格式"""
    def __init__(self, model_name: str = "Qwen/Qwen3-Reranker-8B", instruction: str = ""):
        self.model_name = model_name
        self.instruction = instruction
        self.reranker = None
        self._initialize_reranker()

    def _initialize_reranker(self):
        """初始化Qwen3 Reranker实例"""
        print(f"初始化Qwen3 Reranker: {self.model_name}")
        self.reranker = Qwen3Reranker(self.model_name, self.instruction)
        print("Qwen3 Reranker初始化成功")

    def rerank(self, query: str, nodes: List[NodeWithScore], top_n: int) -> List[NodeWithScore]:
        """对LlamaIndex NodeWithScore列表进行重排序，返回Top-N节点"""
        # 边界条件处理：reranker未初始化或节点为空时，直接返回前N个节点
        if not self.reranker or not nodes:
            return nodes[:top_n]

        # 构建查询-文档对（提取节点文本）
        pairs = [(query, node.node.text) for node in nodes]
        
        # 计算相关性分数
        scores = self.reranker.compute_score(pairs)
        
        # 更新节点分数并按分数降序排序
        for i, node in enumerate(nodes):
            node.score = float(scores[i])
        reranked_nodes = sorted(nodes, key=lambda x: x.score, reverse=True)
        return reranked_nodes[:top_n]

class SimpleVectorStore:
    """基于FAISS的向量存储实现"""
    def __init__(
        self, 
        nodes: List[BaseNode], 
        embed_model, 
        index_path: Optional[str] = None  # 索引文件路径（可选）
    ):
        self.nodes = nodes  # 节点列表（需与索引向量一一对应）
        self.embed_model = embed_model  # 嵌入模型
        self.faiss_index = None
        self.index_path = index_path

        # 优先加载已有索引，否则构建新索引
        if index_path and os.path.exists(index_path):
            self._load_index()
        else:
            self._build_index()
            # 若指定了路径，构建后自动保存
            if index_path:
                self.save_index()

    def _build_index(self):
        """构建新的FAISS索引"""
        print("构建FAISS向量索引...")
        batch_size = 8
        texts = [node.text for node in self.nodes]
        all_embeddings = []

        # 批量生成嵌入
        for i in tqdm(range(0, len(texts), batch_size), desc="生成节点嵌入"):
            batch_texts = texts[i:i+batch_size]
            batch_embeddings = self.embed_model.get_text_embeddings(batch_texts)
            all_embeddings.extend(batch_embeddings)

        embeddings_array = np.array(all_embeddings, dtype=np.float32)
        embedding_dim = embeddings_array.shape[1]
        faiss.normalize_L2(embeddings_array)

        # 初始化扁平内积索引（适合余弦相似度）
        self.faiss_index = faiss.IndexFlatIP(embedding_dim)
        self.faiss_index.add(embeddings_array)
        print(f"✓ 新索引构建完成，包含 {self.faiss_index.ntotal} 个向量（维度：{embedding_dim}）")

    def save_index(self, path: Optional[str] = None):
        """保存索引到本地文件"""
        if not self.faiss_index:
            raise ValueError("索引未初始化，无法保存")
        
        # 优先使用传入的路径，否则使用初始化时的路径
        save_path = path or self.index_path
        if not save_path:
            raise ValueError("请指定索引保存路径")
        
        # 创建父目录（如果不存在）
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        faiss.write_index(self.faiss_index, save_path)
        print(f"✓ 索引已保存至：{save_path}")

    def _load_index(self):
        """从本地文件加载索引"""
        if not self.index_path or not os.path.exists(self.index_path):
            raise FileNotFoundError(f"索引文件不存在：{self.index_path}")
        
        self.faiss_index = faiss.read_index(self.index_path)
        # 验证索引向量数量与节点数量是否匹配
        if self.faiss_index.ntotal != len(self.nodes):
            raise ValueError(
                f"索引向量数量（{self.faiss_index.ntotal}）与节点数量（{len(self.nodes)}）不匹配，可能是索引文件版本不符"
            )
        print(f" 已从 {self.index_path} 加载索引，包含 {self.faiss_index.ntotal} 个向量")

    def search(self, query: str, top_k: int = 10) -> List[NodeWithScore]:
        """检索相似节点"""
        if not self.faiss_index:
            raise ValueError("索引未初始化，无法检索")
        
        # 生成查询嵌入并归一化
        query_embedding = self.embed_model.get_query_embedding(query)
        query_array = np.array([query_embedding], dtype=np.float32)
        faiss.normalize_L2(query_array)

        # FAISS检索（返回分数和索引）
        similarities, top_indices = self.faiss_index.search(query_array, top_k)

        # 封装结果（仅保留正分数）
        results = []
        for idx, score in zip(top_indices[0], similarities[0]):
            if score > 0:
                results.append(NodeWithScore(
                    node=self.nodes[idx],
                    score=float(score)
                ))
        return results

class EnhancedHybridRetriever:
    """增强版混合检索器 - 整合所有优化"""
    
    def __init__(
        self,
        vector_store,
        nodes: List[BaseNode],
        reranker=None,
        query_rewriter=None,
        query_decomposer=None,
        hypo_answer_generator=None,
        deduplicator=None,
        similarity_top_k: int = 100,
        bm25_top_k: int = 100,
        bm25_path: Optional[str] = None,
        bm25_enabled: bool = True
    ):
        self._vector_store = vector_store
        self._nodes = nodes
        self._reranker = reranker
        self._query_rewriter = query_rewriter
        self._query_decomposer = query_decomposer
        self._hypo_answer_generator = hypo_answer_generator
        self._deduplicator = deduplicator
        self._similarity_top_k = similarity_top_k
        self._bm25_top_k = bm25_top_k
        self._bm25_enabled = bm25_enabled
        self._bm25 = None
        self._bm25_path = bm25_path
        self._key_term_extractor = KeyTermExtractor()
        self._rrf = ReciprocalRankFusion()
        
        # 加载BM25索引（如果启用）
        if self._bm25_enabled:
            if bm25_path and os.path.exists(bm25_path):
                self._load_bm25_index()
            else:
                self._build_bm25_index()
                if bm25_path:
                    self._save_bm25_index()
        else:
            print("BM25检索已禁用")

    def _build_bm25_index(self):
        """构建BM25索引"""
        if not self._bm25_enabled:
            return
            
        print("构建BM25索引...")
        corpus = [
            [token for token in jieba.cut(node.text) if len(token.strip()) > 1]
            for node in self._nodes
        ]
        self._bm25 = BM25Okapi(corpus)
        print(f"✓ BM25索引构建完成（{len(corpus)}个节点）")

    def _save_bm25_index(self, path: Optional[str] = None):
        """保存BM25索引"""
        if not self._bm25_enabled or not self._bm25:
            return
            
        import pickle
        save_path = path or self._bm25_path
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        with open(save_path, "wb") as f:
            pickle.dump(self._bm25, f)
        print(f"✓ BM25索引已保存至：{save_path}")

    def _load_bm25_index(self, path: Optional[str] = None):
        """加载BM25索引"""
        if not self._bm25_enabled:
            return
            
        import pickle
        load_path = path or self._bm25_path
        with open(load_path, "rb") as f:
            self._bm25 = pickle.load(f)
        print(f"✓ 已从 {load_path} 加载BM25索引")

    def _single_query_retrieve(self, query: str) -> List[NodeWithScore]:
        """单查询检索"""
        # FAISS向量检索
        vector_nodes = self._vector_store.search(query, self._similarity_top_k)
        vector_node_dict = {n.node.node_id: n.score for n in vector_nodes}

        # BM25检索（如果启用）
        bm25_node_dict = {}
        if self._bm25_enabled and self._bm25:
            query_tokens = [token for token in jieba.cut(query) if len(token.strip()) > 1]
            bm25_scores = self._bm25.get_scores(query_tokens)
            bm25_top_indices = np.argsort(bm25_scores)[::-1][:self._bm25_top_k]
            bm25_node_dict = {
                self._nodes[idx].node_id: float(bm25_scores[idx])
                for idx in bm25_top_indices
                if bm25_scores[idx] > 0}

        # 结果融合
        all_node_ids = set(vector_node_dict.keys()) | set(bm25_node_dict.keys())
        merged_nodes = []

        # 如果有BM25结果，进行融合；否则只使用向量检索结果
        if bm25_node_dict:
            max_bm25_score = max(bm25_node_dict.values()) if bm25_node_dict else 1.0
            for node_id in all_node_ids:
                node = next((n for n in self._nodes if n.node_id == node_id), None)
                if not node:
                    continue

                vector_score = vector_node_dict.get(node_id, 0.0)
                bm25_score = bm25_node_dict.get(node_id, 0.0) / (max_bm25_score + 1e-6)
                combined_score = 0.5 * vector_score + 0.5 * bm25_score
                merged_nodes.append(NodeWithScore(node=node, score=combined_score))
        else:
            # 只使用向量检索
            for node_id, score in vector_node_dict.items():
                node = next((n for n in self._nodes if n.node_id == node_id), None)
                if node:
                    merged_nodes.append(NodeWithScore(node=node, score=score))

        merged_nodes.sort(key=lambda x: x.score, reverse=True)
        return merged_nodes

    def _pre_filter_nodes(self, query: str, nodes: List[NodeWithScore]) -> List[NodeWithScore]:
        """关键信息预过滤 - 基于文档名称计算的智能评分融合"""
        key_terms = self._key_term_extractor.extract_key_terms(query)
        
        if not any(key_terms.values()):
            return nodes
        
        print(f"\n=== 文档名称关键信息智能评分融合 ===")
        print(f"查询: {query}")
        print(f"提取的关键信息: {key_terms}")

        filtered_nodes = []
        adjustment_stats = []
        
        for node in nodes:
            file_name = node.node.metadata.get('file_name', '')
            original_score = node.score

            # 计算关键信息匹配度
            total_boost, detailed_scores = self._key_term_extractor.calculate_key_term_boost(
                file_name, key_terms, query
            )
            
            # 智能评分融合策略
            new_score = self._intelligent_score_fusion(
                original_score, total_boost, detailed_scores, file_name
            )
            
            node.score = new_score
            filtered_nodes.append(node)
            
            # 记录调整统计
            adjustment_stats.append({
                'file_name': file_name,
                'original_score': original_score,
                'new_score': new_score,
                'total_boost': total_boost,
                'adjustment': new_score - original_score
            })
            
            # 打印详细调整信息
            self._print_adjustment_details(
                file_name, original_score, new_score, total_boost, detailed_scores
            )
        
        # 重新排序并输出统计
        filtered_nodes.sort(key=lambda x: x.score, reverse=True)
        self._print_fusion_statistics(adjustment_stats, filtered_nodes)
        
        return filtered_nodes

    def _intelligent_score_fusion(self, original_score: float, total_boost: float, 
                                detailed_scores: Dict[str, float], file_name: str) -> float:
        """智能评分融合策略"""
        
        # 基础调整：使用sigmoid函数进行平滑调整
        def sigmoid_adjustment(x):
            return 4 / (1 + math.exp(-2 * x)) - 2  # 输出范围[-2, 2]
        
        # 1. 对于高原始分数，调整要更谨慎
        if original_score > 1.5:
            adjustment_factor = 0.3  # 高分数文档调整幅度较小
        elif original_score > 0.8:
            adjustment_factor = 0.6  # 中等分数文档调整幅度中等
        else:
            adjustment_factor = 0.8  # 低分数文档调整幅度较大

        # 2. 使用sigmoid函数平滑调整值
        smoothed_boost = sigmoid_adjustment(total_boost) * adjustment_factor
        
        # 3. 考虑各项得分的权重分布
        positive_components = sum([
            detailed_scores['semantic_similarity'],
            detailed_scores['year_match'],
            detailed_scores['policy_match'],
            detailed_scores['location_match'],
            detailed_scores['phrase_match'],
            detailed_scores['time_range_match']
        ])
        
        negative_components = sum([
            detailed_scores['location_penalty'],
            detailed_scores['time_penalty']
        ])
        
        # 4. 计算置信度权重
        confidence_weight = min(1.0, positive_components / 3.0)  # 正项得分越高，置信度越高
        
        # 5. 最终调整计算
        if total_boost > 0:
            # 正向调整：基于置信度加权
            final_adjustment = smoothed_boost * confidence_weight
            new_score = original_score * (1 + final_adjustment)
        else:
            # 负向调整：更谨慎，考虑原始分数
            penalty_severity = 0.5 if original_score > 0.6 else 0.3
            final_adjustment = smoothed_boost * penalty_severity
            new_score = original_score * (1 + final_adjustment)
        
        # 确保分数在合理范围内 [0, 1]
        new_score = max(0.0, min(1.0, new_score))
        
        return new_score

    def _print_adjustment_details(self, file_name: str, original_score: float, 
                                new_score: float, total_boost: float, 
                                detailed_scores: Dict[str, float]):
        """打印调整详情"""
        print(f"\n📄 文档: {file_name}")
        print(f"   原始分数: {original_score:.3f} → 新分数: {new_score:.3f}")
        print(f"   总调整值: {total_boost:.2f}")
        
        # 打印正项得分
        positive_scores = {k: v for k, v in detailed_scores.items() if v > 0 and 'penalty' not in k}
        if positive_scores:
            print("   ✅ 奖励项:")
            for key, value in positive_scores.items():
                if value > 0:
                    print(f"     - {key}: +{value:.2f}")
        
        # 打印负项得分
        negative_scores = {k: v for k, v in detailed_scores.items() if v > 0 and 'penalty' in k}
        if negative_scores:
            print("   ❌ 惩罚项:")
            for key, value in negative_scores.items():
                if value > 0:
                    print(f"     - {key}: -{value:.2f}")

    def _print_fusion_statistics(self, adjustment_stats: List[dict], filtered_nodes: List[NodeWithScore]):
        """打印融合统计信息"""
        print(f"\n=== 智能评分融合统计 ===")
        print(f"处理节点总数: {len(adjustment_stats)}")
        
        if adjustment_stats:
            avg_original = sum(stat['original_score'] for stat in adjustment_stats) / len(adjustment_stats)
            avg_new = sum(stat['new_score'] for stat in adjustment_stats) / len(adjustment_stats)
            avg_adjustment = sum(stat['adjustment'] for stat in adjustment_stats) / len(adjustment_stats)
            
            print(f"平均原始分数: {avg_original:.3f}")
            print(f"平均新分数: {avg_new:.3f}")
            print(f"平均调整幅度: {avg_adjustment:+.3f}")
            
            # 调整幅度分布
            positive_adjustments = [stat for stat in adjustment_stats if stat['adjustment'] > 0]
            negative_adjustments = [stat for stat in adjustment_stats if stat['adjustment'] < 0]
            no_adjustments = [stat for stat in adjustment_stats if stat['adjustment'] == 0]
            
            print(f"分数提升文档: {len(positive_adjustments)}个")
            print(f"分数降低文档: {len(negative_adjustments)}个")
            print(f"分数不变文档: {len(no_adjustments)}个")
        
        if filtered_nodes:
            print(f"调整后最高分: {filtered_nodes[0].score:.3f}")
            print(f"调整后最低分: {filtered_nodes[-1].score:.3f}")

    def retrieve(self, original_query: str) -> List[NodeWithScore]:
        """增强检索主流程"""
        all_queries = [original_query]

        # 1. 查询改写（如果启用）
        if self._query_rewriter and EnhancedConfig.QUERY_REWRITE_ENABLED:
            rewritten_queries = self._query_rewriter.rewrite_queries(
                original_query, EnhancedConfig.QUERY_REWRITE_NUM
            )
            all_queries.extend(rewritten_queries)
        
        # 2. 查询分解（如果启用）
        sub_queries = []
        if self._query_decomposer and EnhancedConfig.QUERY_DECOMPOSE_ENABLED:
            sub_queries = self._query_decomposer.decompose_query(original_query)
            all_queries.extend(sub_queries)

        # 3. 假设答案检索（如果启用）
        if self._hypo_answer_generator and EnhancedConfig.HYPO_ANSWER_ENABLED:
            hypo_answer = self._hypo_answer_generator.generate_hypothetical_answer(original_query)
            all_queries.append(hypo_answer)
        
        # 去重查询
        all_queries = list(set(all_queries))
        print(f"执行多查询检索: {len(all_queries)}个查询")
        print('查询:',all_queries)
        
        # 4. 并行执行多查询检索
        all_ranked_lists = []
        with ThreadPoolExecutor() as executor:
            future_to_query = {
                executor.submit(self._single_query_retrieve, query): query 
                for query in all_queries
            }

            for future in as_completed(future_to_query):
                query = future_to_query[future]
                try:
                    results = future.result()
                    # 关键信息过滤
                    filtered_results = self._pre_filter_nodes(query, results)
                    all_ranked_lists.append(filtered_results)
                except Exception as e:
                    print(f"查询 '{query}' 检索失败: {e}")

        # 5. RRF融合
        if self._rrf and EnhancedConfig.RRF_ENABLED:
            
            if len(all_ranked_lists) > 1:
                fused_nodes = self._rrf.fuse(all_ranked_lists)
            else:
                fused_nodes = all_ranked_lists[0] if all_ranked_lists else []
        
        # 5.5 应用关键信息调整
        # fused_nodes = KeyTermExtractor.apply_key_term_adjustment(fused_nodes, original_query)
        
        # 6. 去重（如果启用）
        if self._deduplicator and EnhancedConfig.DEDUPLICATE_ENABLED:
            deduplicated_nodes = self._deduplicator.deduplicate_nodes(fused_nodes)
        else:
            deduplicated_nodes = fused_nodes

        # 7. 重排序（如果启用）
        if self._reranker and EnhancedConfig.RERANKER_ENABLED and deduplicated_nodes:
            reranked_nodes = self._reranker.rerank(
                original_query,
                deduplicated_nodes[:EnhancedConfig.RERANK_TOP_N],
                EnhancedConfig.RERANK_TOP_N
            )
            final_nodes = reranked_nodes + deduplicated_nodes[EnhancedConfig.RERANK_TOP_N:]
        else:
            final_nodes = deduplicated_nodes
        
        return final_nodes[:EnhancedConfig.FINAL_TOP_K]

class FullyOptimizedQASystem:
    """完全优化的问答系统 - 支持知识图谱数据"""
    
    def __init__(self):
        self.vector_store = None
        self.retriever = None
        self.llm = None
        self.embed_model = None
        self.nodes = []
        self.doc_chunk_map = {}
        self.faiss_path = '../data/output/faiss_index.bin'
        self.bm25_path = '../data/output/bm25_index.pkl'
        self.knowledge_graph_path = '../data/扬州市人工智能产业图谱.xlsx'
        self.include_knowledge_graph = None

        # 组件实例
        self.query_rewriter = None
        self.query_decomposer = None
        self.hypo_answer_generator = None
        self.reranker = None
        self.deduplicator = None
        self.splitter = None

    def initialize(self, doc_path: str, include_knowledge_graph: bool = True):
        """初始化系统 - 可选择是否包含知识图谱数据
        
        Args:
            doc_path: 传统文档路径
            include_knowledge_graph: 是否包含知识图谱数据，默认为True
        """
        self.include_knowledge_graph = EnhancedConfig.include_knowledge_graph
        
        # 1. 加载模型
        print("\n1. 加载模型...")
        self.embed_model = Qwen3Embedding(
            model_name=EnhancedConfig.EMBEDDING_MODEL,
            instruction=EnhancedConfig.EMBEDDING_INSTRUCTION)

        # 2. 处理文档（可选择是否包含知识图谱）
        print("\n2. 读取和处理文档...")
        processor = EnhancedDocumentProcessor()
        documents = processor.read_all_documents(
            doc_path, 
            knowledge_graph_path=self.knowledge_graph_path,
            include_knowledge_graph=self.include_knowledge_graph
        )

        # 3. 文档分割
        print(f"\n3. 文档分割 - 使用 {EnhancedConfig.CHUNK_MODE} 模式...")
        self.splitter = DocumentSplitterFactory.create_splitter()
        
        all_chunks = []
        for doc in documents:
            chunks = self.splitter.split_document(doc)
            all_chunks.extend(chunks)
        
        self.nodes = all_chunks
        print(f" 生成 {len(self.nodes)} 个{EnhancedConfig.CHUNK_MODE}节点")
        
        # 统计不同类型文档的数量
        source_types = {}
        for node in self.nodes:
            source_type = node.metadata.get('source_type', 'unknown')
            source_types[source_type] = source_types.get(source_type, 0) + 1
        
        print("文档类型统计:")
        for source_type, count in source_types.items():
            print(f"  - {source_type}: {count}")

        # 构建文档-段落映射
        for node in self.nodes:
            file_name = node.metadata.get('file_name', '未知文件')
            if file_name not in self.doc_chunk_map:
                self.doc_chunk_map[file_name] = []
            self.doc_chunk_map[file_name].append(node)

        # 4. 构建索引
        print("\n4. 构建向量索引...")
        self.vector_store = SimpleVectorStore(
            nodes=self.nodes,
            embed_model=self.embed_model,
            index_path=self.faiss_path
        )

        # 5. 加载重排序器
        if EnhancedConfig.RERANKER_ENABLED:
            print("\n5. 加载重排序器...")
            self.reranker = Qwen3DirectReranker(
                model_name=EnhancedConfig.RERANKER_MODEL,
                instruction=EnhancedConfig.RERANKER_INSTRUCTION)
        else:
            print("\n5. 重排序器已禁用")
            self.reranker = None
        
        # 6. 加载LLM和组件
        print("\n6. 加载LLM和组件...")
        self.llm = SimpleLLM(model_path=EnhancedConfig.GENERATION_MODEL)
        
        # 初始化可选组件
        if EnhancedConfig.QUERY_REWRITE_ENABLED:
            self.query_rewriter = QueryRewriter(self.llm)
            print("  - 查询改写器已启用")
        else:
            self.query_rewriter = None
            print("  - 查询改写器已禁用")
            
        if EnhancedConfig.QUERY_DECOMPOSE_ENABLED:
            self.query_decomposer = QueryDecomposer(self.llm)
            print("  - 查询分解器已启用")
        else:
            self.query_decomposer = None
            print("  - 查询分解器已禁用")
            
        if EnhancedConfig.HYPO_ANSWER_ENABLED:
            self.hypo_answer_generator = HypotheticalAnswerGenerator(self.llm)
            print("  - 假设答案生成器已启用")
        else:
            self.hypo_answer_generator = None
            print("  - 假设答案生成器已禁用")
            
        if EnhancedConfig.DEDUPLICATE_ENABLED:
            self.deduplicator = SimHashDeduplicator(EnhancedConfig)
            print("  - 去重器已启用")
        else:
            self.deduplicator = None
            print("  - 去重器已禁用")

        # 7. 创建增强检索器
        print("\n7. 初始化增强检索器...")
        self.retriever = EnhancedHybridRetriever(
            vector_store=self.vector_store,
            nodes=self.nodes,
            reranker=self.reranker,
            query_rewriter=self.query_rewriter,
            query_decomposer=self.query_decomposer,
            hypo_answer_generator=self.hypo_answer_generator,
            deduplicator=self.deduplicator,
            bm25_path=self.bm25_path,
            bm25_enabled=EnhancedConfig.BM25_ENABLED
        )

        print("\n✓ 系统初始化完成！")
        print(f"  知识图谱集成: {'已加载' if EnhancedConfig.include_knowledge_graph else '已禁用'}")
        print(f"  总文档节点: {len(self.nodes)}")
        print(f"  组件状态:")
        print(f"  - 查询改写: {'启用' if EnhancedConfig.QUERY_REWRITE_ENABLED else '禁用'}")
        print(f"  - 查询分解: {'启用' if EnhancedConfig.QUERY_DECOMPOSE_ENABLED else '禁用'}")
        print(f"  - 假设答案: {'启用' if EnhancedConfig.HYPO_ANSWER_ENABLED else '禁用'}")
        print(f"  - 去重: {'启用' if EnhancedConfig.DEDUPLICATE_ENABLED else '禁用'}")
        print(f"  - 重排序: {'启用' if EnhancedConfig.RERANKER_ENABLED else '禁用'}")
        print(f"  - BM25: {'启用' if EnhancedConfig.BM25_ENABLED else '禁用'}")

    def get_adjacent_chunks(self, current_node, n=1):
        """获取相邻块"""
        file_name = current_node.metadata.get('file_name', '未知文件')
        
        if file_name not in self.doc_chunk_map:
            return [], []
        
        chunks = self.doc_chunk_map[file_name]
        
        # 根据分割模式选择ID字段
        if EnhancedConfig.CHUNK_MODE == 'paragraph':
            current_id = current_node.metadata.get('paragraph_id', -1)
        else:
            current_id = current_node.metadata.get('chunk_id', -1)
            
        if current_id < 0:
            return [], []
        
        total_chunks = len(chunks)
        
        prev_chunks = []
        next_chunks = []
        
        # 前序块
        start_prev = max(0, current_id - n)
        for idx in range(current_id - 1, start_prev - 1, -1):
            prev_chunks.append(chunks[idx])

        # 后序块
        end_next = min(total_chunks - 1, current_id + n)
        for idx in range(current_id + 1, end_next + 1):
            next_chunks.append(chunks[idx])
        
        return prev_chunks, next_chunks

    def _build_context_str(self, retrieved_nodes, n=1):
        """构建上下文 - 集成SimHash去重确保无重复段落"""
        context_parts = []
        total_length = 0
        max_context_length = EnhancedConfig.max_context_length
        
        # 去重相关集合
        added_chunk_ids = set()
        added_content_hashes = set()  # 用于存储已添加内容的SimHash
        
        # 统计信息
        total_retrieved = len(retrieved_nodes)
        low_score_skipped = 0
        duplicate_skipped = 0
        simhash_duplicate_skipped = 0
        added_count = 0
        threshold = 0.2
        
        print(f"\n=== 上下文构建统计 ===")
        print(f"检索到的节点总数: {total_retrieved}")
        print(f"分数阈值: {threshold}")
        print(f"相邻块数量: {n}")
        print(f"最大上下文长度: {max_context_length}")
        print(f"SimHash相似阈值: {EnhancedConfig.SIMILAR_THRESHOLD}")
        
        # 首先对检索到的节点按分数排序（降序）
        sorted_nodes = sorted(retrieved_nodes, key=lambda x: x.score, reverse=True)
        
        for node_with_score in sorted_nodes:
            if node_with_score.score <= threshold:
                low_score_skipped += 1
                continue

            node = node_with_score.node
            file_name = node.metadata.get('file_name', '未知文件')

            # 根据分割模式选择ID字段
            if EnhancedConfig.CHUNK_MODE == 'paragraph':
                chunk_id = node.metadata.get('paragraph_id', '未知')
            else:
                chunk_id = node.metadata.get('chunk_id', '未知')

            chunk_key = f"{file_name}_{chunk_id}"
            
            if chunk_key in added_chunk_ids:
                duplicate_skipped += 1
                continue

            # 获取相邻块
            prev_chunks, next_chunks = self.get_adjacent_chunks(node, n)

            # 构建当前节点及其相邻块的内容
            chunks_to_add = []
            
            # 添加前序相邻块
            for prev_chunk in prev_chunks:
                if EnhancedConfig.CHUNK_MODE == 'paragraph':
                    prev_chunk_id = prev_chunk.metadata.get('paragraph_id', '未知')
                else:
                    prev_chunk_id = prev_chunk.metadata.get('chunk_id', '未知')
                prev_key = f"{file_name}_{prev_chunk_id}"
                if prev_key not in added_chunk_ids:
                    chunks_to_add.append((prev_chunk, prev_key, f"前序-{prev_chunk_id}"))
            
            # 添加当前块
            chunks_to_add.append((node, chunk_key, f"当前-{chunk_id}"))
            
            # 添加后续相邻块
            for next_chunk in next_chunks:
                if EnhancedConfig.CHUNK_MODE == 'paragraph':
                    next_chunk_id = next_chunk.metadata.get('paragraph_id', '未知')
                else:
                    next_chunk_id = next_chunk.metadata.get('chunk_id', '未知')
                next_key = f"{file_name}_{next_chunk_id}"
                if next_key not in added_chunk_ids:
                    chunks_to_add.append((next_chunk, next_key, f"后续-{next_chunk_id}"))

            # 处理所有要添加的块，进行SimHash去重
            for chunk, chunk_key, chunk_type in chunks_to_add:
                if chunk_key in added_chunk_ids:
                    continue
                    
                chunk_text = chunk.text.strip()
                if len(chunk_text) < EnhancedConfig.MIN_TEXT_LENGTH:
                    # 对于过短的文本，直接添加
                    chunk_info = f"【{chunk_type}】得分：{node_with_score.score:.2f} 文档: {file_name}\n{chunk_text}\n\n"
                    
                    if total_length + len(chunk_info) > max_context_length:
                        break
                    
                    context_parts.append(chunk_info)
                    total_length += len(chunk_info)
                    added_chunk_ids.add(chunk_key)
                    added_count += 1
                    print(f"添加短文本节点: {file_name} ({chunk_type}), 分数: {node_with_score.score:.3f}, 当前总长度: {total_length}")
                    continue
                
                # 计算当前块的SimHash
                chunk_hash = self.deduplicator._simhash(chunk_text)
                
                # 检查是否与已添加内容重复
                is_duplicate = False
                for existing_hash in added_content_hashes:
                    # print('文档距离：', self.deduplicator._hamming_distance(chunk_hash, existing_hash))
                    
                    if self.deduplicator._hamming_distance(chunk_hash, existing_hash) <= EnhancedConfig.SIMILAR_THRESHOLD:
                        
                        is_duplicate = True
                        simhash_duplicate_skipped += 1
                        print(f"SimHash去重: {file_name} ({chunk_type}), 检测到重复内容")
                        break
                
                if not is_duplicate:
                    chunk_info = f"【{chunk_type}】得分：{node_with_score.score:.2f} 文档: {file_name}\n{chunk_text}\n\n"
                    
                    if total_length + len(chunk_info) > max_context_length:
                        print(f"达到最大上下文长度限制，停止添加更多内容")
                        break
                    
                    context_parts.append(chunk_info)
                    total_length += len(chunk_info)
                    added_chunk_ids.add(chunk_key)
                    added_content_hashes.add(chunk_hash)
                    added_count += 1
                    print(f"添加节点: {file_name} ({chunk_type}), 分数: {node_with_score.score:.3f}, 当前总长度: {total_length}")
            
            if total_length >= max_context_length:
                print(f"达到最大上下文长度限制，停止处理更多节点")
                break
        
        # 输出统计信息
        print(f"\n=== 去重统计结果 ===")
        print(f"总检索节点数: {total_retrieved}")
        print(f"低分跳过: {low_score_skipped}")
        print(f"重复ID跳过: {duplicate_skipped}")
        print(f"SimHash重复跳过: {simhash_duplicate_skipped}")
        print(f"成功添加: {added_count}")
        print(f"最终上下文长度: {total_length}")
        
        return "".join(context_parts)

    def build_prompt(self, question: str, context_str: str) -> str:
        """构建提示词 - 支持知识图谱数据"""
        return f"""
# 角色定位
您是一位资深的扬州市政策咨询与人工智能产业链招商专家，具备深厚的政策解读能力和产业分析经验。

# 任务指令
请基于以下提供的政策文档和产业图谱数据，对用户问题进行详细、准确、全面的解答。

# 可用数据源
{context_str}

# 待回答问题
{question}

# 回答要求
## 准确性要求
- 严格依据提供的政策文档和产业图谱数据，所有信息必须与原文保持一致
- 关键信息（数字、日期、政策条款、企业名称、业务范围、地区等）必须完整复刻原文
- 严禁任何形式的杜撰、改写、概括或主观推断

## 内容深度要求
- 对政策条款进行完整解读，说明政策适用范围、条件、标准等关键要素
- 对产业信息进行全面分析，包括企业分布、业务特点、区域特征等
- 针对列举类问题，必须完整呈现所有相关条目，不得遗漏任何重要信息
- 对于复杂问题，需要从多个维度进行深入分析，确保回答的全面性

## 结构化要求
- 回答内容应该层次清晰，逻辑严谨
- 重要信息可以适当使用分点（1.2.3...）说明，但避免使用特殊符号
- 保持语句的自然流畅，同时确保信息的完整传达

## 专业性要求
- 使用政策文档和产业图谱中的专业术语和核心概念
- 保持政策语言的严谨性和产业分析的深度
- 在准确的基础上，确保回答易于理解
- 直接回答问题，不用复述题目

## 完整性要求
- 回答必须完全覆盖用户问题的所有方面
- 对于涉及多个政策或多个企业的问题，需要整合所有相关信息
- 确保回答能够完整地解决用户问题

请基于上述要求，提供专业、准确、全面的回答。
"""

    def answer(self, question: str) -> str:
        """问答接口"""
        if not self.retriever:
            return "系统未初始化"
        
        # 检索相关文档
        retrieved_nodes = self.retriever.retrieve(question)
        if not retrieved_nodes:
            return "未找到相关政策信息。"
        
        # 构建上下文
        context_str = self._build_context_str(retrieved_nodes, n=1)
        print(context_str)
        # 生成答案
        prompt = self.build_prompt(question, context_str)
        messages = [ChatMessage(role="user", content=prompt)]
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = self.llm.chat(messages)
                answer = response.content
                break
            except Exception as e:
                if attempt == max_retries - 1:
                    return f"生成答案时出错: {str(e)}"
                time.sleep(2)
        
        # 后处理
        return self._post_process(answer)

    def _post_process(self, answer: str) -> str:
        """答案后处理"""
        # 清理格式
        pattern = r'(.*?)</think>(.*)'
        match = re.search(pattern, answer, re.DOTALL)
        
        if match:
            content = match.group(2).strip()
        else:
            content = answer.strip()
        
        final_answer = re.sub(r'\n+', '\n', content).strip()
        return final_answer